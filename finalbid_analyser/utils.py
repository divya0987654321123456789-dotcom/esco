# # utils.py
import os
import json
import torch
import pandas as pd
from doctr.models import ocr_predictor
from doctr.io import DocumentFile
import requests
from docx import Document as DocxDocument
import re
from typing import List, Dict, Any
import io
from functools import lru_cache
try:
    from sentence_transformers import SentenceTransformer, util as st_util
except Exception:
    SentenceTransformer, st_util = None, None

# ----------------- OCR & Text Extraction -------------------
@torch.no_grad()
def load_doctr_model(device: str = "cuda"):
    try:
        if not torch.cuda.is_available():
            device = "cpu"
        return ocr_predictor(pretrained=True).to(device)
    except Exception:
        return ocr_predictor(pretrained=True).to("cpu")

def extract_text_with_doctr(pdf_bytes: bytes, model) -> str:
    doc = DocumentFile.from_pdf(pdf_bytes)
    result = model(doc)
    return "\n\n--- Page Break ---\n\n".join(page.render() for page in result.pages)

# Prefer text extraction via pypdf when available; fallback to OCR
def extract_text_with_pypdf(pdf_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(pdf_bytes))
        texts: List[str] = []
        for page in reader.pages:
            try:
                texts.append(page.extract_text() or "")
            except Exception:
                texts.append("")
        return "\n\n--- Page Break ---\n\n".join([t.strip() for t in texts if t is not None])
    except Exception as e:
        return ""

def extract_text_best(pdf_bytes: bytes, doctr_model) -> str:
    # Try direct text extraction first
    text = extract_text_with_pypdf(pdf_bytes)
    if text and len(text) >= 500:
        return text
    # Fallback to OCR
    try:
        return extract_text_with_doctr(pdf_bytes, doctr_model)
    except Exception:
        return text or ""

# ----------------- OpenAI-Compatible Client -------------------
def generate_json_via_api(prompt: str, timeout_s: int = 120) -> dict:
    api_key = os.getenv("OPENAI_API_KEY") or os.getenv("GROQ_API_KEY")
    base_url = os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1/chat/completions")
    model = os.getenv("OPENAI_MODEL", "gpt-4o")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY (or GROQ_API_KEY) is required to run without Ollama.")
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    body = {
        "model": model,
        "temperature": 0,
        "response_format": {"type": "json_object"},
        "messages": [{"role": "user", "content": prompt}],
    }
    try:
        resp = requests.post(base_url, headers=headers, json=body, timeout=timeout_s)
        resp.raise_for_status()
        content = resp.json()["choices"][0]["message"]["content"]
        return json.loads(content)
    except Exception as e:
        print(f"API request failed: {e}")
        return {}

# ----------------- File Reader -------------------
def read_company_file(file_path: str, doctr_model) -> str:
    try:
        with open(file_path, 'rb') as file:
            if file_path.lower().endswith(".txt"):
                return file.read().decode("utf-8", errors="ignore")
            elif file_path.lower().endswith(".docx"):
                doc = DocxDocument(file)
                return "\n".join([p.text for p in doc.paragraphs])
            elif file_path.lower().endswith(".xlsx"):
                df = pd.read_excel(file)
                return df.astype(str).to_string()
            elif file_path.lower().endswith(".pdf"):
                pdf_bytes = file.read()
                return extract_text_best(pdf_bytes, doctr_model)
            else:
                return file.read().decode("utf-8", errors="ignore")
    except Exception as e:
        print(f"Error reading company file {file_path}: {e}")
        return ""

# ----------------- Local Pre-Weighting (Fallback Logic) -------------------
def apply_scope_preweights(rfp_info: dict, full_text: str) -> dict:
    LIGHTING_KEYWORDS = ["lighting", "led", "illumination", "fixture", "luminaire"]
    PROJECT_TYPE_KEYWORDS = {
        "hvac": {"IKIO": 0, "": 10, "IKIO Energy": 10}, "solar": {"IKIO": 0, "": 10, "IKIO Energy": 10},
        "water": {"IKIO": 0, "": 10, "IKIO Energy": 10}, "waste water": {"IKIO": 0, "": 10, "IKIO Energy": 10},
        "wastewater": {"IKIO": 0, "": 10, "IKIO Energy": 10}, "building envelope": {"IKIO": 0, "": 10, "IKIO Energy": 10},
        "esco": {"IKIO": 0, "": 10, "IKIO Energy": 10}, "energy saving": {"IKIO": 0, "": 10, "IKIO Energy": 10},
        "generator": {"IKIO": 0, "": 10, "IKIO Energy": 10},
    }
    filename_part = rfp_info.get('filename', '').lower()
    text = (filename_part + " " + full_text).lower()
    preweights = {"IKIO": 0, "": 0, "IKIO Energy": 0}
    detected_type, condition, detection_source = "Unknown", "N/A", "N/A"

    if any(k in text for k in LIGHTING_KEYWORDS):
        detected_type = "Lighting"
        has_supply = rfp_info.get('scope_supply', False)
        has_install = rfp_info.get('scope_installation', False)
        substitution_allowed = rfp_info.get('scope_substitution_allowed', False)
        no_substitution = rfp_info.get('scope_no_substitution', False)
        detection_source = "Filename + Content" if any(kw in filename_part for kw in LIGHTING_KEYWORDS) else "Content"

        if has_supply and not has_install and substitution_allowed:
            preweights, condition = {"IKIO": 10, "": 0, "IKIO Energy": 0}, "Supply + Substitution Allowed"
        elif has_supply and has_install and substitution_allowed:
            preweights, condition = {"IKIO": 10, "": 10, "IKIO Energy": 10}, "Supply + Installation + Substitution Allowed"
        elif has_supply and has_install and no_substitution:
            preweights, condition = {"IKIO": 0, "": 10, "IKIO Energy": 10}, "Supply + Installation + Substitution Not Allowed"
        elif not has_supply and has_install:
            preweights, condition = {"IKIO": 0, "": 10, "IKIO Energy": 10}, "Installation Only"
        else:
            preweights, condition = {"IKIO": 10, "": 0, "IKIO Energy": 0}, "Lighting (Default All)"
    else:
        for kw, weights in PROJECT_TYPE_KEYWORDS.items():
            if kw in text:
                detected_type, preweights, condition = kw.title(), weights, f"Detected via keyword '{kw}'"
                detection_source = "Filename + Content" if kw in filename_part else "Content"
                break
    return {"detected_type": detected_type, "condition": condition, "preweights": preweights, "detection_source": detection_source}

# ----------------- Criteria 1 via API (Primary Logic) -------------------
def apply_scope_preweights_via_api(rfp_summary: dict, rfp_filename: str | None = None, raw_text: str | None = None) -> dict:
    api_key = os.getenv("OPENAI_API_KEY")
    base_url = os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1/chat/completions")
    model = os.getenv("OPENAI_MODEL", "gpt-4o")
    compact = (f"project_type: {rfp_summary.get('project_type','')}; scope: {rfp_summary.get('scope','')}; summary: {rfp_summary.get('summary','')}; ")[:6000]
    lt = (raw_text or compact).lower()
    local_has_supply = any(k in lt for k in ["supply", "furnish", "provide", "procure", "deliver", "material"])
    local_has_install = any(k in lt for k in ["install", "installation", "installing", "installed", "installment", "retrofit", "replacement", "replace", "erection"])
    local_sub_allowed = any(k in lt for k in ["substitution allowed", "approved equal", "or equal", "or equivalent", "approved equivalent"])
    local_no_sub = any(k in lt for k in ["no substitution", "no substitutions", "no alternates", "no equals", "as specified only"])
    
    fallback_info = {
        'filename': rfp_filename or '', 'scope_supply': local_has_supply, 'scope_installation': local_has_install,
        'scope_substitution_allowed': local_sub_allowed, 'scope_no_substitution': local_no_sub, 'project_type': rfp_summary.get('project_type','')
    }

    if not api_key:
        return apply_scope_preweights(fallback_info, compact)

    prompt = f"""
Return ONLY JSON with fields: {{"detected_type": "Lighting|HVAC|...", "has_supply": true|false, "has_install": true|false, "substitution_allowed": true|false, "no_substitution": true|false}}
Use the filename (if provided) and this compact RFP summary to classify type and scope.
FILENAME: {rfp_filename or 'N/A'}
TEXT: {compact}"""
    try:
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        body = {"model": model, "temperature": 0, "response_format": {"type": "json_object"}, "messages": [{"role": "user", "content": prompt}]}
        resp = requests.post(base_url, headers=headers, json=body, timeout=60)
        resp.raise_for_status()
        content = resp.json()["choices"][0]["message"]["content"]
        data = json.loads(content)
        dtype = str(data.get("detected_type", "Unknown")).strip()
        has_supply = bool(data.get("has_supply", False)) or local_has_supply
        has_install = bool(data.get("has_install", False)) or local_has_install
        sub_allowed = bool(data.get("substitution_allowed", False)) or local_sub_allowed
        no_sub = bool(data.get("no_substitution", False)) or local_no_sub

        if dtype.lower() == "lighting":
            if has_supply and not has_install and sub_allowed: preweights, condition = {"IKIO": 10, "": 0, "IKIO Energy": 0}, "Supply + Substitution Allowed"
            elif has_supply and has_install and sub_allowed: preweights, condition = {"IKIO": 10, "": 10, "IKIO Energy": 10}, "Supply + Installation + Substitution Allowed"
            elif has_supply and has_install and no_sub: preweights, condition = {"IKIO": 0, "": 10, "IKIO Energy": 10}, "Supply + Installation + Substitution Not Allowed"
            elif not has_supply and has_install: preweights, condition = {"IKIO": 0, "": 10, "IKIO Energy": 10}, "Installation Only"
            else: preweights, condition = {"IKIO": 10, "": 0, "IKIO Energy": 0}, "Lighting (Default All)"
            return {"detected_type": "Lighting", "condition": condition, "preweights": preweights, "detection_source": "API"}
        else:
            mapping = { "hvac": {"IKIO": 0, "": 10, "IKIO Energy": 10}, "solar": {"IKIO": 0, "": 10, "IKIO Energy": 10}, "water": {"IKIO": 0, "": 10, "IKIO Energy": 10}, "waste water": {"IKIO": 0, "": 10, "IKIO Energy": 10}, "building envelope": {"IKIO": 0, "": 10, "IKIO Energy": 10}, "esco": {"IKIO": 0, "": 10, "IKIO Energy": 10}, "energy saving": {"IKIO": 0, "": 10, "IKIO Energy": 10}, "generator": {"IKIO": 0, "": 10, "IKIO Energy": 10} }
            key = dtype.lower()
            if key in mapping:
                return {"detected_type": dtype.title(), "condition": f"Detected via API as {dtype}", "preweights": mapping[key], "detection_source": "API"}
    except Exception as e:
        print(f"API call for scope pre-weighting failed: {e}. Falling back to local rules.")
    
    return apply_scope_preweights(fallback_info, compact)

# ----------------- Criteria 2: Project Location -------------------
def apply_location_preweights(raw_text: str, rfp_summary: dict, rfp_filename: str | None = None) -> dict:
    text_full = "\n".join([rfp_filename or "", rfp_summary.get("scope", "") or "", rfp_summary.get("summary", "") or "", raw_text or ""])
    t = text_full.lower()
    anchors = ["project location", "project is located", "site location", "location:", "city:", "state:", "located at", "located in", "place of performance", "place of work", "project address", "site address"]
    windows = [t[max(0, idx-100): idx+200] for a in anchors for idx in [i for i, _ in enumerate(t) if t.startswith(a, i)]]
    if not windows: windows = [t[:3000]]
    
    in_patterns = [r"\bindiana\b", r"\bindianapolis\b", r"\bindianpolis\b", r"indian\-polis"]
    tx_patterns = [r"\bdallas\b", r"\btexas\b", r"\btx\b"]
    exclude_patterns = ["registered", "licens", "vendor", "bidder", "mailing", "office", "headquarter"]

    def window_valid(win: str) -> bool: return not any(ex in win for ex in exclude_patterns)
    def any_regex(patterns: list, hay: str) -> bool: return any(re.search(p, hay) for p in patterns)

    evidence_snippets = [re.sub(r"\s+", " ", w.strip())[:240] for w in windows if window_valid(w) and (any_regex(in_patterns, w) or any_regex(tx_patterns, w))]
    found_indiana = any(window_valid(w) and any_regex(in_patterns, w) for w in windows)
    found_dallas = any(window_valid(w) and any_regex(tx_patterns, w) for w in windows)
    
    pre = {"IKIO": 0, "": 0, "IKIO Energy": 0}
    matched, location_names = [], []
    if found_indiana:
        pre["IKIO"], pre["IKIO Energy"] = 10, 10
        matched.append("Indiana/Indianapolis")
        location_names.append("Indianapolis, IN" if any_regex([r"\bindianapolis\b", r"\bindianpolis\b"], t) else "Indiana")
    if found_dallas:
        pre[""] = 10
        matched.append("Dallas/Texas")
        location_names.append("Dallas, TX" if any_regex([r"\bdallas\b"], t) else "Texas")

    return {"preweights": pre, "detected_locations": location_names or ["Unknown"], "scored_locations": matched, "location_display": " / ".join(location_names) if location_names else "Unknown", "detection_source": "Anchored Location Detection", "evidence_snippets": evidence_snippets[:5]}

# ----------------- Detailed Checklist Evaluation -------------------
def load_checklist_items(excel_path: str = "checklist.xlsx") -> List[dict]:
    try:
        df = pd.read_excel(excel_path)
        cols = {c.strip().lower(): c for c in df.columns}
        req_col = cols.get("requirement") or cols.get("requirements") or list(df.columns)[0]
        crit_col = cols.get("criticality")
        items = []
        for _, row in df.iterrows():
            requirement = str(row.get(req_col, "")).strip()
            if not requirement: continue
            criticality = str(row.get(crit_col, "C")).strip().upper() if crit_col else "C"
            items.append({"requirement": requirement, "criticality": criticality or "C"})
        return items
    except Exception: return [{"requirement": "Scope alignment with RFP", "criticality": "C"}, {"requirement": "Schedule feasibility", "criticality": "C"}]

def build_compact_rfp_context(rfp_summary: dict, max_chars: int = 3000) -> str:
    parts = [f"Project Type: {rfp_summary.get('project_type','')}"]
    if rfp_summary.get("scope"): parts.append(f"Scope: {rfp_summary.get('scope','')}")
    if rfp_summary.get("summary"): parts.append(f"Summary: {rfp_summary.get('summary','')}")
    for key in ["key_requirements", "technical_requirements", "bid_requirements"]:
        vals = rfp_summary.get(key) or []
        if isinstance(vals, list) and vals: parts.append(f"{key}: {', '.join(map(str, vals[:20]))}")
    return "\n".join(parts)[:max_chars]

def find_best_snippet(text: str, requirement: str, max_len: int = 180) -> str:
    if not text or not requirement: return ""
    try:
        cleaned = re.sub(r"\s+", " ", str(text))
        parts = re.split(r"(?<=[.!?])\s+|\n+|;\s+", cleaned)
        stop = {"the", "a", "an", "and", "or", "of", "to", "for", "in", "on", "at", "by", "with", "is", "are", "be", "as", "that", "this", "these", "those", "from"}
        req_tokens = [t for t in re.findall(r"[a-z0-9]+", requirement.lower()) if t not in stop] or re.findall(r"[a-z0-9]+", requirement.lower())
        best_score, best = 0, ""
        for p in parts:
            ptoks = [t for t in re.findall(r"[a-z0-9]+", p.lower()) if t not in stop]
            if not ptoks: continue
            overlap = len(set(req_tokens) & set(ptoks))
            if overlap > best_score:
                best_score, best = overlap, p.strip()
        return (best or cleaned)[:max_len]
    except Exception: return str(text)[:max_len]

# ----------------- Semantic Reranker -------------------
@lru_cache(maxsize=1)
def load_st_model(model_name: str = "all-MiniLM-L6-v2"):
    try:
        if SentenceTransformer is None:
            return None
        return SentenceTransformer(model_name)
    except Exception as e:
        print(f"Failed to load sentence-transformer model '{model_name}': {e}")
        return None

def split_into_spans(text: str, max_span_chars: int = 280) -> List[str]:
    if not text:
        return []
    try:
        # Split by sentence-ish boundaries and aggregate into spans of similar length
        parts = re.split(r"(?<=[.!?])\s+|\n+|;\s+", re.sub(r"\s+", " ", str(text)))
        spans: List[str] = []
        current = ""
        for p in parts:
            if not p.strip():
                continue
            if len(current) + 1 + len(p) <= max_span_chars:
                current = f"{current} {p}".strip()
            else:
                if current:
                    spans.append(current.strip())
                current = p
        if current:
            spans.append(current.strip())
        return spans
    except Exception:
        return [str(text)[:max_span_chars]]

def rerank_snippets(text: str, query: str, top_k: int = 2, model_name: str = "all-MiniLM-L6-v2") -> List[str]:
    model = load_st_model(model_name)
    if model is None or st_util is None:
        # Fallback to heuristic if model unavailable
        best = find_best_snippet(text, query, max_len=220)
        return [best] if best else []
    spans = split_into_spans(text)
    if not spans:
        return []
    try:
        emb_spans = model.encode(spans, convert_to_tensor=True, normalize_embeddings=True)
        emb_query = model.encode([query], convert_to_tensor=True, normalize_embeddings=True)
        sims = st_util.cos_sim(emb_query, emb_spans)[0]
        # Get top indices
        k = max(1, min(int(top_k), len(spans)))
        topk = torch.topk(sims, k)
        idxs = topk.indices.tolist()
        # Keep original order by similarity rank
        ranked = [spans[i] for i in idxs]
        return ranked
    except Exception as e:
        print(f"Reranker failed: {e}")
        best = find_best_snippet(text, query, max_len=220)
        return [best] if best else []

def evaluate_checklist_with_api(
    company_name: str,
    rfp_summary: dict,
    company_text: str,
    items: List[dict],
    rfp_full_text: str = "",
    mode: str = "combined",
    temperature: float = 0.0,
    use_reranker: bool = False,
    reranker_top_k: int = 2,
    reranker_model_name: str | None = None,
) -> List[dict]:
    api_key = os.getenv("OPENAI_API_KEY")
    base_url = os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1/chat/completions")
    model = os.getenv("OPENAI_MODEL", "gpt-4o")

    rfp_ctx = build_compact_rfp_context(rfp_summary)
    company_ctx = (company_text or "")[:4000]

    use_rfp = mode in ("combined", "rfp_only")
    use_company = mode in ("combined", "company_only")

    def make_hint(requirement: str) -> dict:
        return {
            "requirement": requirement,
            "criticality": "C",
            "rfp_hint": find_best_snippet(rfp_full_text if use_rfp else "", requirement),
            "company_hint": find_best_snippet(company_text if use_company else "", requirement),
        }

    items_text = []
    for it in items:
        req = it["requirement"]
        if use_reranker:
            rfp_hits = rerank_snippets(rfp_full_text if use_rfp else "", req, top_k=reranker_top_k or 2, model_name=reranker_model_name or "all-MiniLM-L6-v2")
            comp_hits = rerank_snippets(company_text if use_company else "", req, top_k=reranker_top_k or 2, model_name=reranker_model_name or "all-MiniLM-L6-v2")
            rfp_hint = " | ".join(rfp_hits)[:240]
            company_hint = " | ".join(comp_hits)[:240]
        else:
            rfp_hint = find_best_snippet(rfp_full_text if use_rfp else "", req)
            company_hint = find_best_snippet(company_text if use_company else "", req)
        items_text.append({
            "requirement": req,
            "criticality": it.get("criticality", "C"),
            "rfp_hint": rfp_hint,
            "company_hint": company_hint,
        })

    mode_instructions = {
        "rfp_only": (
            "Evaluate only against the RFP context. Ignore company documents. "
            "Return rfp_context; set company_context to empty."
        ),
        "company_only": (
            "Evaluate only against the company documents. Ignore the RFP context. "
            "Return company_context; set rfp_context to empty."
        ),
        "combined": (
            "Evaluate using both RFP context and company documents. "
            "Return both rfp_context and company_context."
        ),
    }

    prompt = (
        "Return ONLY JSON with array 'results': {requirement, criticality, evaluation, suggestions, rfp_context, company_context, comments}.\n"
        "- evaluation: Yes|Partial|No.\n"
        "- rfp_context: evidence from RFP (<=25 words).\n"
        "- company_context: evidence from company doc (<=25 words).\n"
        "- suggestions: action/observation (<=20 words).\n\n"
        f"Mode: {mode}. {mode_instructions.get(mode, '')}\n\n"
        f"RFP Context:\n{rfp_ctx if use_rfp else ''}\n\n"
        f"Company Context:\n{company_ctx if use_company else ''}\n\n"
        f"Checklist with HINTS:\n{json.dumps(items_text, ensure_ascii=False)}\n\n"
        "JSON schema: {\"results\":[{\"requirement\":\"...\",...}]}"
    )
    try:
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        body = {
            "model": model,
            "temperature": max(0.0, min(1.0, float(temperature))),
            "response_format": {"type": "json_object"},
            "messages": [{"role": "user", "content": prompt}],
        }
        resp = requests.post(base_url, headers=headers, json=body, timeout=90)
        resp.raise_for_status()
        data = resp.json()["choices"][0]["message"]["content"]
        parsed = json.loads(data).get("results", [])
        results: List[dict] = []
        for it in parsed:
            evaluation_value = it.get("evaluation", "Partial")
            score_value = 2 if evaluation_value == "Yes" else 1 if evaluation_value == "Partial" else 0
            results.append(
                {
                    "requirement": it.get("requirement", ""),
                    "criticality": it.get("criticality", "C"),
                    "evaluation": evaluation_value,
                    "suggestions": it.get("suggestions", ""),
                    "rfp_context": it.get("rfp_context", "") if use_rfp else "",
                    "company_context": it.get("company_context", "") if use_company else "",
                    "comments": it.get("comments", ""),
                    "score": score_value,
                }
            )
        return results
    except Exception as e:
        print(f"Checklist evaluation API call failed for {company_name} ({mode}): {e}")
        return [
            {
                "requirement": it["requirement"],
                "criticality": it.get("criticality", "C"),
                "evaluation": "Partial",
                "suggestions": "Evaluation failed",
                "rfp_context": "" if not use_rfp else find_best_snippet(rfp_full_text, it["requirement"]),
                "company_context": "" if not use_company else find_best_snippet(company_text, it["requirement"]),
                "comments": "",
                "score": 1,
            }
            for it in items
        ]

def evaluate_checklist_with_consensus(
    company_name: str,
    rfp_summary: dict,
    company_text: str,
    items: List[dict],
    rfp_full_text: str = "",
    mode: str = "combined",
    votes: int = 1,
    temperature: float = 0.2,
    use_reranker: bool = False,
    reranker_top_k: int = 2,
    reranker_model_name: str | None = None,
) -> List[dict]:
    try:
        total_votes = max(1, min(5, int(votes)))
    except Exception:
        total_votes = 1

    if total_votes == 1:
        return evaluate_checklist_with_api(
            company_name,
            rfp_summary,
            company_text,
            items,
            rfp_full_text,
            mode,
            temperature=0.0,
            use_reranker=use_reranker,
            reranker_top_k=reranker_top_k,
            reranker_model_name=reranker_model_name,
        )

    runs: List[List[dict]] = []
    for _ in range(total_votes):
        runs.append(
            evaluate_checklist_with_api(
                company_name,
                rfp_summary,
                company_text,
                items,
                rfp_full_text,
                mode,
                temperature=temperature,
                use_reranker=use_reranker,
                reranker_top_k=reranker_top_k,
                reranker_model_name=reranker_model_name,
            )
        )

    # Aggregate by requirement text
    requirement_to_results: Dict[str, List[dict]] = {}
    for run in runs:
        for res in run:
            req = str(res.get("requirement", "")).strip()
            if not req:
                continue
            requirement_to_results.setdefault(req, []).append(res)

    aggregated: List[dict] = []
    for base in items:
        req = str(base.get("requirement", "")).strip()
        if not req:
            continue
        votes_for_req = requirement_to_results.get(req, [])
        if not votes_for_req:
            aggregated.append({
                "requirement": req,
                "criticality": base.get("criticality", "C"),
                "evaluation": "Partial",
                "suggestions": "Insufficient evidence",
                "rfp_context": "",
                "company_context": "",
                "comments": "",
                "score": 1,
                "confidence": 0.0,
            })
            continue

        # Majority vote on evaluation
        eval_counts: Dict[str, int] = {}
        for r in votes_for_req:
            key = str(r.get("evaluation", "Partial"))
            eval_counts[key] = eval_counts.get(key, 0) + 1
        majority_eval = sorted(eval_counts.items(), key=lambda kv: kv[1], reverse=True)[0][0]
        agreement = eval_counts.get(majority_eval, 0)
        confidence = round(agreement / total_votes, 3)
        score_value = 2 if majority_eval == "Yes" else 1 if majority_eval == "Partial" else 0

        # Choose contexts/suggestions from the run(s) that match the majority eval
        majority_runs = [r for r in votes_for_req if str(r.get("evaluation", "Partial")) == majority_eval]
        picked = majority_runs[0] if majority_runs else votes_for_req[0]

        aggregated.append({
            "requirement": req,
            "criticality": picked.get("criticality", base.get("criticality", "C")),
            "evaluation": majority_eval,
            "suggestions": picked.get("suggestions", ""),
            "rfp_context": picked.get("rfp_context", ""),
            "company_context": picked.get("company_context", ""),
            "comments": picked.get("comments", ""),
            "score": score_value,
            "confidence": confidence,
        })

    return aggregated

# ----------------- Simple Company Context Cache -------------------
def get_company_cache_path(company: str) -> str:
    cache_dir = "company_context_cache"
    os.makedirs(cache_dir, exist_ok=True)
    return os.path.join(cache_dir, f"{company}.txt")

def save_company_context(company: str, text: str) -> None:
    try:
        with open(get_company_cache_path(company), "w", encoding="utf-8") as f: f.write(text or "")
    except Exception: pass

def load_company_context(company: str) -> str:
    try:
        with open(get_company_cache_path(company), "r", encoding="utf-8") as f: return f.read()
    except Exception: return ""

# ----------------- DOCX Export -------------------
def export_to_docx(state: Dict[str, Any], file_path: str):
    doc = DocxDocument()
    doc.add_heading("Bid Compliance Evaluation Report", level=1)
    doc.add_heading("RFP Information", level=2)
    doc.add_paragraph(f"Filename: {state.get('rfp_filename', 'N/A')}")
    doc.add_paragraph(f"Project Type: {state.get('rfp_summary', {}).get('project_type', 'N/A')}")
    scope_res = state.get('preweight_scope_results', {})
    loc_res = state.get('preweight_location_results', {})
    combined = state.get('combined_preweights', {})
    doc.add_heading("Pre-Weighting Analysis", level=2)
    doc.add_paragraph(f"Scope/Type Detection: {scope_res.get('detected_type', 'N/A')} ({scope_res.get('condition', 'N/A')})")
    doc.add_paragraph(f"Location Detection: {loc_res.get('location_display', 'N/A')}")
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text, hdr_cells[3].text = 'Company', 'Scope Score', 'Location Score', 'Total Score'
    for company in ["IKIO", "IKIO Energy"]:
        row_cells = table.add_row().cells
        row_cells[0].text = company
        row_cells[1].text = f"{scope_res.get('preweights', {}).get(company, 0)}/10"
        row_cells[2].text = f"{loc_res.get('preweights', {}).get(company, 0)}/10"
        row_cells[3].text = f"{combined.get(company, 0)}/20"
    doc.add_paragraph(f"\nTop companies (by pre-weight): {', '.join(state.get('top_companies', []))}")

    # Render evaluations: support both legacy single-set and new multi-set structure
    for name, eval_data in state.get('company_evaluations', {}).items():
        doc.add_page_break()
        doc.add_heading(f"Detailed Evaluation: {name}", level=2)

        def add_eval_section(title: str, data: Dict[str, Any]):
            doc.add_heading(title, level=3)
            fs_local = data.get("finalScoring", {})
            doc.add_paragraph(
                f"Recommendation: {fs_local.get('recommendation', 'N/A')} "
                f"({fs_local.get('overallCompliancePercentage', 0)}% Compliance)"
            )
            eval_table = doc.add_table(rows=1, cols=4)
            h = eval_table.rows[0].cells
            h[0].text, h[1].text, h[2].text, h[3].text = "Requirement", "Evaluation", "Context", "Suggestions"
            for item in data.get('complianceAssessment', []):
                r = eval_table.add_row().cells
                conf = item.get('confidence')
                conf_str = f" ({int(round(conf*100))}% conf)" if isinstance(conf, (int, float)) else ""
                r[0].text = item.get('requirement', '')
                r[1].text = f"{item.get('evaluation', '')}{conf_str}"
                r[2].text = (
                    (f"RFP: {item.get('rfp_context', '')}\n" if item.get('rfp_context') else "") +
                    (f"Company: {item.get('company_context', '')}" if item.get('company_context') else "")
                ).strip()
                r[3].text = item.get('suggestions', '')

        # New structure with three sections
        if any(k in eval_data for k in ("rfp_only", "company_only", "combined")):
            if "rfp_only" in eval_data:
                add_eval_section("RFP Checklist (RFP-only)", eval_data["rfp_only"])
            if "company_only" in eval_data:
                add_eval_section("Company Checklist (Company-only)", eval_data["company_only"])
            if "combined" in eval_data:
                add_eval_section("Combined Checklist (RFP + Company)", eval_data["combined"])
        else:
            # Legacy structure
            add_eval_section("Combined Checklist (Legacy)", eval_data)

    doc.save(file_path)
