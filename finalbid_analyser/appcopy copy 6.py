# -*- coding: utf-8 -*-

import os, json, re, tempfile, time, threading, hashlib, uuid, mimetypes
import streamlit as st
from openai import OpenAI
from io import BytesIO
from PIL import Image
import pymysql
from contextlib import nullcontext
from dotenv import load_dotenv
from typing import List, Dict, Optional
from datetime import date, datetime, timedelta
import logging, asyncio

# RAG Model Support (fallback when OpenAI API unavailable)
_RAG_AVAILABLE = False
_RAG_MODEL = None
try:
    from sentence_transformers import SentenceTransformer
    _RAG_AVAILABLE = True
except Exception:
    _RAG_AVAILABLE = False
try:
    import requests
    from bs4 import BeautifulSoup
    _SCRAPE_OK = True
except Exception:
    _SCRAPE_OK = False


# ---- Simple RAG retrieval over Ikio DB context (embeddings) ----
# Uses SentenceTransformer when available. This is used to find evidence snippets when exact keywords are missing.
try:
    import numpy as np
    _NP_OK = True
except Exception:
    _NP_OK = False

@st.cache_resource
def _get_embedder():
    if not _RAG_AVAILABLE or not _NP_OK:
        return None
    try:
        return SentenceTransformer("all-MiniLM-L6-v2")
    except Exception:
        return None

def _chunk_text(text: str, chunk_size: int = 900, overlap: int = 150) -> list[str]:
    if not text:
        return []
    text = re.sub(r"\s+", " ", str(text)).strip()
    if not text:
        return []
    chunks: list[str] = []
    i = 0
    n = len(text)
    while i < n:
        j = min(n, i + chunk_size)
        chunks.append(text[i:j])
        # advance with overlap
        i = max(i + chunk_size - overlap, i + 1)
    return chunks

def _cosine_sim(a: "np.ndarray", b: "np.ndarray") -> float:
    denom = (float(np.linalg.norm(a)) * float(np.linalg.norm(b))) + 1e-9
    return float(np.dot(a, b) / denom)

def rag_retrieve(company_blob: str, query: str, top_k: int = 3) -> list[dict]:
    """Return top_k snippets: [{score: float, text: str}, ...]."""
    embedder = _get_embedder()
    if embedder is None or not company_blob or not query:
        return []

    chunks = _chunk_text(company_blob)
    if not chunks:
        return []

    try:
        qv = np.array(embedder.encode([query])[0], dtype=np.float32)
        cvs = np.array(embedder.encode(chunks), dtype=np.float32)
        scored = []
        for idx, cv in enumerate(cvs):
            scored.append((idx, _cosine_sim(qv, cv)))
        scored.sort(key=lambda x: x[1], reverse=True)
        out = []
        for idx, s in scored[:max(1, int(top_k))]:
            out.append({"score": round(float(s), 4), "text": chunks[idx]})
        return out
    except Exception:
        return []

# ---- OCR (DocTR) ----
try:
    from doctr.io import DocumentFile
    from doctr.models import ocr_predictor
    _DOCTR_OK = True
except Exception:
    _DOCTR_OK = False

# ---- DOCX ----
try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

# Optional fast PDF text extraction (avoid OCR when PDF already contains selectable text)
try:
    from pypdf import PdfReader as _PdfReader  # type: ignore
except Exception:
    try:
        from PyPDF2 import PdfReader as _PdfReader  # type: ignore
    except Exception:
        _PdfReader = None

# Optional Excel parsing (batch/folder evaluation)
try:
    import pandas as _pd  # type: ignore
except Exception:
    _pd = None

def extract_text_from_pdf_bytes_fast(file_bytes: bytes, max_pages: int = 30) -> str:
    """Extract selectable text from PDFs quickly (no OCR). Returns empty string if unavailable."""
    if not file_bytes or _PdfReader is None:
        return ""
    try:
        reader = _PdfReader(BytesIO(file_bytes))
        parts: list[str] = []
        n = 0
        for page in getattr(reader, "pages", []) or []:
            n += 1
            if max_pages and n > int(max_pages):
                break
            try:
                t = page.extract_text() or ""
            except Exception:
                t = ""
            t = re.sub(r"[ \t]+\n", "\n", (t or ""))
            t = re.sub(r"\n{3,}", "\n\n", t)
            if t.strip():
                parts.append(t.strip())
        return "\n\n".join(parts).strip()
    except Exception:
        return ""

def extract_text_from_excel_bytes(file_bytes: bytes, filename: str = "") -> str:
    """Best-effort extraction of text from Excel files (.xlsx/.xls)."""
    if not file_bytes:
        return ""
    if _pd is None:
        return ""
    try:
        bio = BytesIO(file_bytes)
        # Read all sheets, stringify cells, join rows
        sheets = _pd.read_excel(bio, sheet_name=None, header=None)  # type: ignore[arg-type]
        parts: list[str] = []
        for sheet_name, df in (sheets or {}).items():
            try:
                df = df.fillna("")
                # stringify all cells
                df = df.astype(str)
                rows = ["\t".join([c for c in row if c and c.strip()]) for row in df.values.tolist()]
                rows = [r for r in rows if r.strip()]
                if rows:
                    parts.append(f"=== SHEET: {sheet_name} ===\n" + "\n".join(rows))
            except Exception:
                continue
        return ("\n\n".join(parts)).strip()
    except Exception:
        return ""

def extract_text_any(
    file_bytes: bytes,
    filename: str,
    *,
    prefer_fast_pdf: bool = True,
    force_pdf_ocr: bool = False,
    pdf_fast_max_pages: int = 30,
) -> str:
    """Unified text extraction for supported formats."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        if not force_pdf_ocr:
            fast = extract_text_from_pdf_bytes_fast(file_bytes, max_pages=pdf_fast_max_pages)
            # If PDF has embedded/selectable text, this will be non-trivial in length.
            if prefer_fast_pdf and fast and len(fast.strip()) >= 300:
                return fast
        # Fallback to OCR (slow) for scanned/image PDFs
        return extract_text(file_bytes, filename)
    if name.endswith((".xlsx", ".xls")):
        return extract_text_from_excel_bytes(file_bytes, filename)
    return extract_text(file_bytes, filename)

# ----------------- STREAMLIT & ENV -----------------
try:
    st.set_page_config(page_title="RFP Evaluation System", layout="wide")
except Exception:
    # When imported by other Streamlit pages, set_page_config may have been called already
    pass
load_dotenv()

# Reduce noisy disconnect logs from Tornado/asyncio when client disconnects
def _install_log_filters():
    try:
        logging.getLogger("tornado.websocket").setLevel(logging.CRITICAL)
        logging.getLogger("tornado.application").setLevel(logging.ERROR)
        logging.getLogger("tornado.general").setLevel(logging.ERROR)
        logging.getLogger("tornado.access").setLevel(logging.WARNING)
        logging.getLogger("asyncio").setLevel(logging.CRITICAL)
    except Exception:
        pass
    try:
        def _loop_exception_handler(loop, context):
            exc = context.get("exception")
            if exc is not None:
                name = exc.__class__.__name__
                if name in ("WebSocketClosedError", "StreamClosedError"):
                    return
            msg = (context.get("message") or "").lower()
            if ("websocketclosederror" in msg) or ("stream is closed" in msg):
                return
            loop.default_exception_handler(context)
        loop = asyncio.get_event_loop()
        loop.set_exception_handler(_loop_exception_handler)
    except Exception:
        pass
_install_log_filters()

# --------------- OneDrive scan logging ---------------
def _get_onedrive_log_path() -> str:
    env_path = os.getenv("ONEDRIVE_SCAN_LOG", "").strip()
    if env_path:
        return env_path
    return os.path.join(APP_ROOT if "APP_ROOT" in globals() else os.getcwd(), "onedrive_scan.log")

def _get_onedrive_logger() -> logging.Logger:
    logger = logging.getLogger("onedrive_scan")
    if getattr(logger, "_ikio_initialized", False):
        return logger
    logger.setLevel(logging.INFO)
    logger.propagate = False
    log_path = _get_onedrive_log_path()
    try:
        os.makedirs(os.path.dirname(log_path), exist_ok=True)
    except Exception:
        pass
    try:
        fh = logging.FileHandler(log_path, encoding="utf-8")
        fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s")
        fh.setFormatter(fmt)
        logger.addHandler(fh)
    except Exception:
        pass
    logger._ikio_initialized = True  # type: ignore[attr-defined]
    return logger

# --------------- Streamlit-safe UI helpers ---------------
try:
    from streamlit.runtime.scriptrunner import get_script_run_ctx as _st_get_ctx
except Exception:
    _st_get_ctx = None

def _can_update_ui() -> bool:
    try:
        return (_st_get_ctx is not None) and (_st_get_ctx() is not None)
    except Exception:
        return False

def _st_safe(func, *args, **kwargs):
    if not _can_update_ui():
        return None
    try:
        return func(*args, **kwargs)
    except Exception:
        return None

def _st_spinner(message: str):
    # Returns a real spinner if UI is active, else a no-op context
    if _can_update_ui():
        return st.spinner(message)
    return nullcontext()

def _get_openai_api_key() -> str:
    try:
        secret_key = (st.secrets.get("OPENAI_API_KEY", "").strip())
    except Exception:
        secret_key = ""
    env_key = os.getenv("OPENAI_API_KEY", "").strip()
    return env_key or secret_key

def _get_openai_client():
    key = _get_openai_api_key()
    if not key:
        return None
    try:
        return OpenAI(api_key=key)
    except Exception:
        return None

DEFAULT_MODEL = "gpt-4o"

def _get_selected_model() -> str:
    try:
        sel = st.session_state.get("llm_model")
    except Exception:
        sel = None
    return (sel or DEFAULT_MODEL)

# ----------------- DB CONFIG -----------------
def _select_db_name() -> str:
    desired = (os.getenv("MYSQL_DB", "esco_v23_clean") or "").strip() or "esco_v23_clean"
    fallback = (os.getenv("MYSQL_DB_FALLBACK", "esco_v23_clean") or "").strip() or "esco_v23_clean"
    if desired == "esco" and fallback:
        return fallback
    return desired

DB_CFG = dict(
    host=os.getenv("MYSQL_HOST", "localhost"),
    user=os.getenv("MYSQL_USER", "root"),
    password=os.getenv("MYSQL_PASSWORD", ""),
    database=_select_db_name(),
    cursorclass=pymysql.cursors.DictCursor,
    # Fail fast if DB is not reachable to avoid long UI hangs
    connect_timeout=int(os.getenv("MYSQL_CONNECT_TIMEOUT", "3")),
)

def _open_mysql_or_create():
    try:
        conn = pymysql.connect(**DB_CFG)
        return True, conn, DB_CFG, None
    except Exception as e:
        try:
            err_txt = str(e)
            if "1813" in err_txt or "Tablespace for table" in err_txt:
                fallback = (os.getenv("MYSQL_DB_FALLBACK", "esco_v23_clean") or "").strip() or "esco_v23_clean"
                if DB_CFG.get("database") != fallback:
                    cfg = dict(DB_CFG)
                    cfg["database"] = fallback
                    conn = pymysql.connect(**cfg)
                    return True, conn, cfg, None
        except Exception:
            pass
        _st_safe(st.error, f"MySQL connection failed: {e}")
        return False, None, DB_CFG, e

# --------------- Constants (aligned with new_app_5.py 512-638) ---------------
APP_ROOT = os.path.dirname(os.path.abspath(__file__))
CACHE_DIR = os.path.join(APP_ROOT, "company_context_cache")
UPLOADS_DIR = os.path.join(APP_ROOT, "uploads")
PROJECT_ROOT = os.path.abspath(os.path.join(APP_ROOT, os.pardir))
BID_DOCS_DIR = os.path.join(PROJECT_ROOT, "static", "bid_documents")
os.makedirs(CACHE_DIR, exist_ok=True)
os.makedirs(UPLOADS_DIR, exist_ok=True)
os.makedirs(BID_DOCS_DIR, exist_ok=True)
COMPANY_LOCATIONS = {
    "Ikio Led Lighting LLC": "Indianapolis, IN",
    "Sunsprint Engineering": "Batesville, IN",
    "METCO Engineering, Inc.": "Dallas, TX",
}

# License states (Ikio Led Lighting LLC list provided by you)
COMPANY_LICENSE_STATES = {
    "Ikio Led Lighting LLC": {
        "AL","AZ","AR","CA","CO","CT","FL","GA","IL","IN","IA",
        "LA","ME","MI","MN","NE","NH","NY","NC","OH","OK",
        "OR","PA","RI","SD","TN","TX","UT","VA","WA","WV","WI"
    },
    # Update these as you get real data:
    "METCO Engineering, Inc.": {"TX","IN","OK","LA","AR"},
    
    "Sunsprint Engineering": {"IN","KY","OH"},
}

COMPANY_WEBSITES = {
    # Use UPPERCASE keys for robust matching
    # NOTE: If your Ikio website differs, set it in the Company Database page; DB value will be used as fallback.
    "IKIO LED LIGHTING LLC": "https://www.ikioledlighting.com",
    "METCO": "https://www.metcoengineering.com",
    "METCO ENGINEERING, INC.": "https://www.metcoengineering.com",
    "SUNSPRINT": "https://www.sunsprintengineering.com",
    "SUNSPRINT ENGINEERING": "https://www.sunsprintengineering.com",
}

def _get_mapped_website(company_name: str) -> str | None:
    n = (company_name or "").strip().upper()
    if not n:
        return None
    # direct match or best-effort containment for minimal aliasing
    if n in COMPANY_WEBSITES:
        return COMPANY_WEBSITES[n]
    for key, url in COMPANY_WEBSITES.items():
        k = (key or "").strip().upper()
        if not k:
            continue
        if n == k or n.endswith(k) or k in n:
            return url
    return None

MAPPING_TEXT = (
    """
    - Lighting (Supply +Subsitution Allowed + Installation of Equivalent or Similar + New LED Lights Installation ) → Ikio Led Lighting LLC
    - Lighting (Supply+Installation+Subsitution Not allowed) → Sunsprint Engineering or METCO Engineering, Inc.
    - Lighting (Supply + Installation +Subsitution Allowed) → Ikio Led Lighting LLC or METCO Engineering, Inc. or Sunsprint Engineering
    - HVAC → Sunsprint Engineering or METCO Engineering, Inc.
    - Solar PV → Sunsprint Engineering or METCO Engineering, Inc.
    - Lighting (Only Installation) → Sunsprint Engineering or METCO Engineering, Inc.
    - Lighting (Installation of Equivalent or Similar) → Ikio Led Lighting LLC
    - Water Management → Sunsprint Engineering or METCO Engineering, Inc.
    - Building Envelope → Sunsprint Engineering or METCO Engineering, Inc.
    - Construction → Sunsprint Engineering or METCO Engineering, Inc.
    - ESCO → Sunsprint Engineering or METCO Engineering, Inc.
    - Emergency Generator → Sunsprint Engineering or METCO Engineering, Inc.
    """
).strip()

EVAL_STEPS_TEXT = (
    """
    GENERALIZED RFP EVALUATION TABLES FOR EPC COMPANIES
    -------------------------------------------------
    Step 1: Extract the following information from the RFP.
      - Submission Method (normalize to one or more of: Electronically; Email; Submission via link; Bids return to Email ID; Sealed Bid; Hand Delivery; Mailed Bid). 
        Look for cues such as: "electronic submission", "submit electronically", "via portal", "online portal", "upload", "email to", "send to email", "submit by email", 
        "submission link", "apply at link", "return to <email>", "sealed bid", "sealed envelope", "hand deliver", "deliver in person", "mail", "mailed bids", "USPS", "FedEx", "UPS".
    Step 2: Write an exact concise summary of the scope of work and key requirements/documents.
    Step 3: Answer the following information from the RFP:
      a. What is the project State?
      b. What kind of license is required to work on the project?
      c. Is any site investigation/visit required/mandatory?
      d. Are any specific procurement requirements (BABA/BAA/Davis Bacon) are there?
      e. What specific qualifications are required for the project?
      f. Does the project require SBE, MBE, WBE, HUB goals?
      g. Does the project require and specific security clearance or working hour restrictions?
      h. Is any bond (payment/performance/bid) is required?
      i. Is any insurance is required?

    COMPANY SELECTION CRITERIA
    --------------------------
    Step 1: Recommend applicable companies per mapping.
    Step 2: Assign 5 base points to each recommended company.
    Step 3: Add 5 bonus points if project state matches company’s state.
    Step 4: Suggest best companies accordingly (if all three companies have equal points, then evaluate scoring for all three companies).
    Step 5 (Override Rule): If the detected work profile is "Lighting (Supply + Installation +Subsitution Allowed)", you MUST evaluate and score ALL THREE companies (Ikio Led Lighting LLC, METCO Engineering, Inc., Sunsprint Engineering) in the BID Evaluation Process, not just top two.

    BID EVALUATION PROCESS
    ----------------------
    Step 1: Answer these questions for the best 2 recommended company with their company’s documents and compare with the RFP’s response, then give remarks and score according to the criteria mentioned:
    a. Is project state and company state same?
    b. Is the company or its subcontractor (if available) has required license as mentioned in the BID document in the project state (if required/mandatory)?
    c. If site investigation/visit required/mandatory, can company/subcontractor do this?
    d. Is company capable of fulfilling specific procurement requirements (BABA/BAA/Davis Bacon) (if required/mandatory)?
    e. Is company capable of fulfilling specific qualifications for the project (if mandatory/required)?
    f. Can companies meet SBE, MBE, WBE, HUB goals (if required)?
    g. Can company meet specific security clearance or working hour restrictions (if required)?
    h. Can company provide bond (payment/performance/bid) (if required)?
    i. Can company provide insurance (if required)?
    Step 2: Produce a BID Evaluation Table listing: Question, Score, Remark, Recommendation.
    Step 3: Compute total score per company.
    Step 4: Determine Go/No-Go per company with rationale.
    Step 5: If scores tie, use qualifications fit for tiebreaker.
    Step 6: Provide final recommendations to qualify for BID.
    """
).strip()

# Rule-based Company Recommendation (Step 1 logic)
PROFILE_TO_COMPANIES = {
    # new_app_5 names
    "Lighting (Supply + Substitution Allowed + New LED Lights Installation + Installation of Equivalent or Similar)": ["Ikio Led Lighting LLC"],
    "Lighting (Supply + Installation + Substitution Not allowed)": ["Sunsprint Engineering", "METCO Engineering, Inc."],
    "Lighting (Supply + Installation + Substitution Allowed)": ["Ikio Led Lighting LLC", "METCO Engineering, Inc.", "Sunsprint Engineering"],
    "HVAC": ["Sunsprint Engineering", "METCO Engineering, Inc."],
    "Solar PV": ["Sunsprint Engineering", "METCO Engineering, Inc."],
    "Lighting (Only Installation)": ["Sunsprint Engineering", "METCO Engineering, Inc."],
    "Lighting (Installation of Equivalent or Similar)": ["Ikio Led Lighting LLC"],
    "Water Management": ["Sunsprint Engineering", "METCO Engineering, Inc."],
    "Building Envelope": ["Sunsprint Engineering", "METCO Engineering, Inc."],
    "Construction": ["Sunsprint Engineering", "METCO Engineering, Inc."],
    "ESCO": ["Sunsprint Engineering", "METCO Engineering, Inc."],
    "Emergency Generator": ["Sunsprint Engineering", "METCO Engineering, Inc."],
    # compatibility with previous profile strings
    "Lighting (Supply + Substitution Allowed)": ["Ikio Led Lighting LLC"],
}

PROFILE_PRIORITY_ORDER = [
    "Lighting (Supply + Installation + Substitution Allowed + New LED Lights Installation)",
    "Lighting (Supply + Installation + Substitution Not allowed)",
    "Lighting (Supply + Substitution Allowed)",
    "Lighting (Only Installation)",
    "Lighting (Installation of Equivalent or Similar)",
    "HVAC",
    "Solar PV",
    "Water Management",
    "Building Envelope",
    "Construction",
    "ESCO",
    "Emergency Generator",
]

# ----------------- TABLES ENSURE / SAVE -----------------
def _ensure_company_tables(conn):
    with conn.cursor() as c:
        c.execute("""
        CREATE TABLE IF NOT EXISTS company_details (
          id INT AUTO_INCREMENT PRIMARY KEY,
          company_name VARCHAR(100) NOT NULL UNIQUE,
          website VARCHAR(255), address VARCHAR(255), start_date DATE,
          created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        ) ENGINE=InnoDB CHARSET=utf8mb4;""")
        c.execute("""
        CREATE TABLE IF NOT EXISTS company_capabilities (
          id INT AUTO_INCREMENT PRIMARY KEY,
          company_name VARCHAR(100) NOT NULL,
          capability_title VARCHAR(255),
          capability_description TEXT,
          naics_codes VARCHAR(255),
          created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
          INDEX(company_name)
        ) ENGINE=InnoDB CHARSET=utf8mb4;""")
        c.execute("""
        CREATE TABLE IF NOT EXISTS company_preferences (
          id INT AUTO_INCREMENT PRIMARY KEY,
          company_name VARCHAR(100) NOT NULL UNIQUE,
          deal_breakers TEXT, deal_makers TEXT,
          federal BOOLEAN DEFAULT TRUE, state_local BOOLEAN DEFAULT TRUE,
          preferred_states VARCHAR(255), preferred_countries VARCHAR(255),
          created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        ) ENGINE=InnoDB CHARSET=utf8mb4;""")
        c.execute("""
        CREATE TABLE IF NOT EXISTS company_locations (
          id INT AUTO_INCREMENT PRIMARY KEY,
          company_name VARCHAR(100) NOT NULL UNIQUE,
          base_location VARCHAR(255),
          base_state VARCHAR(10),
          created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        ) ENGINE=InnoDB CHARSET=utf8mb4;""")
    conn.commit()

def _ensure_scrape_tables(conn):
    with conn.cursor() as c:
        c.execute("""
        CREATE TABLE IF NOT EXISTS company_web_context (
          id INT AUTO_INCREMENT PRIMARY KEY,
          company_name VARCHAR(100) NOT NULL UNIQUE,
          url VARCHAR(512),
          content LONGTEXT,
          last_fetched TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
        ) ENGINE=InnoDB CHARSET=utf8mb4;""")
    conn.commit()

def _normalize_url(url: str | None) -> str | None:
    if not url:
        return None
    u = url.strip()
    if not u:
        return None
    if not re.match(r"^https?://", u, flags=re.IGNORECASE):
        u = "https://" + u
    return u

# ----------------- COMPANY UPSERT API (used by pages/Company_Database.py) -----------------
def upsert_company_details(conn, company_name, website, address, start_date):
    _ensure_company_tables(conn)
    with conn.cursor() as c:
        c.execute(
            (
                "INSERT INTO company_details (company_name, website, address, start_date) "
                "VALUES (%s,%s,%s,%s) "
                "ON DUPLICATE KEY UPDATE website=VALUES(website), address=VALUES(address), start_date=VALUES(start_date)"
            ),
            (company_name, website, address, start_date),
        )
    conn.commit()

def upsert_company_preferences(conn, company_name, deal_breakers, deal_makers, federal, state_local, preferred_states, preferred_countries):
    _ensure_company_tables(conn)
    with conn.cursor() as c:
        c.execute(
            (
                "INSERT INTO company_preferences (company_name, deal_breakers, deal_makers, federal, state_local, preferred_states, preferred_countries) "
                "VALUES (%s,%s,%s,%s,%s,%s,%s) "
                "ON DUPLICATE KEY UPDATE deal_breakers=VALUES(deal_breakers), deal_makers=VALUES(deal_makers), federal=VALUES(federal), state_local=VALUES(state_local), preferred_states=VALUES(preferred_states), preferred_countries=VALUES(preferred_countries)"
            ),
            (company_name, deal_breakers, deal_makers, federal, state_local, preferred_states, preferred_countries),
        )
    conn.commit()

def add_company_capability(conn, company_name, title, desc, naics):
    _ensure_company_tables(conn)
    with conn.cursor() as c:
        c.execute(
            (
                "INSERT INTO company_capabilities (company_name, capability_title, capability_description, naics_codes) "
                "VALUES (%s,%s,%s,%s)"
            ),
            (company_name, title, desc, naics),
        )
    conn.commit()

def upsert_company_location(conn, company_name: str, base_location: str | None, base_state: str | None = None):
    _ensure_company_tables(conn)
    with conn.cursor() as c:
        c.execute(
            (
                "INSERT INTO company_locations (company_name, base_location, base_state) "
                "VALUES (%s,%s,%s) "
                "ON DUPLICATE KEY UPDATE base_location=VALUES(base_location), base_state=VALUES(base_state)"
            ),
            (company_name, base_location, base_state),
        )
    conn.commit()

def _ensure_eval_tables(conn):
    with conn.cursor() as c:
        c.execute("""
        CREATE TABLE IF NOT EXISTS evaluation_runs (
          id INT AUTO_INCREMENT PRIMARY KEY,
          file_name VARCHAR(255),
          file_hash VARCHAR(64) NULL,
          result_text LONGTEXT NOT NULL,
          export_json LONGTEXT NULL,
          created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
          INDEX(file_hash),
          INDEX(file_name)
        ) ENGINE=InnoDB CHARSET=utf8mb4;""")

        # Ensure newer schema on existing DBs (add file_hash if missing)
        try:
            c.execute("ALTER TABLE evaluation_runs ADD COLUMN file_hash VARCHAR(64) NULL")
        except Exception:
            pass
        try:
            c.execute("CREATE INDEX idx_evaluation_runs_file_hash ON evaluation_runs (file_hash)")
        except Exception:
            pass
        c.execute("""
        CREATE TABLE IF NOT EXISTS bid_incoming (
          id INT AUTO_INCREMENT PRIMARY KEY,
          b_name VARCHAR(255),
          due_date VARCHAR(255),
          state VARCHAR(64),
          scope TEXT,
          type VARCHAR(255),
          scoring INT,
          comp_name VARCHAR(255),
          ik_project_code VARCHAR(32),
          decision VARCHAR(16),
          summary LONGTEXT,
          evaluation_date DATE NULL,
          created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        ) ENGINE=InnoDB CHARSET=utf8mb4;""")
        # Ensure newer schema on existing DBs (add evaluation_date if missing)
        try:
            c.execute("ALTER TABLE bid_incoming ADD COLUMN evaluation_date DATE NULL")
        except Exception:
            pass
        try:
            c.execute("ALTER TABLE bid_incoming ADD COLUMN ik_project_code VARCHAR(32) NULL")
        except Exception:
            pass
        # Speed up bid-level duplicate detection (best-effort)
        try:
            c.execute("CREATE INDEX idx_bid_incoming_bname_due_date ON bid_incoming (b_name, due_date)")
        except Exception:
            pass
        c.execute("""
        CREATE TABLE IF NOT EXISTS uploaded_rfp_files (
          id INT AUTO_INCREMENT PRIMARY KEY,
          original_filename VARCHAR(255) NOT NULL,
          saved_filename VARCHAR(255) NOT NULL,
          file_path VARCHAR(512) NOT NULL,
          file_size BIGINT,
          file_hash VARCHAR(64),
          file_type VARCHAR(50),
          uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
          INDEX(original_filename),
          INDEX(uploaded_at),
          INDEX(file_hash)
        ) ENGINE=InnoDB CHARSET=utf8mb4;""")
        # Add file_hash index if it doesn't exist (for faster duplicate detection)
        try:
            c.execute("CREATE INDEX IF NOT EXISTS idx_file_hash ON uploaded_rfp_files(file_hash)")
        except Exception:
            pass
    conn.commit()

def _ensure_bid_assign_meta_table(conn):
    with conn.cursor() as c:
        c.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_assign_meta (
              g_id INT PRIMARY KEY,
              data LONGTEXT,
              updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
            ) ENGINE=InnoDB CHARSET=utf8mb4;
            """
        )
    conn.commit()

def _safe_filename(name: str) -> str:
    safe = re.sub(r"[^\w\-.]+", "_", name or "").strip("_")
    return safe or f"file_{uuid.uuid4().hex}.bin"

def _store_bid_documents(bid_id: int, files: list[tuple[str, bytes]]) -> list[dict]:
    if not bid_id or not files:
        return []
    stored: list[dict] = []
    folder = os.path.join(BID_DOCS_DIR, str(bid_id))
    try:
        os.makedirs(folder, exist_ok=True)
    except Exception:
        return []
    for fname, fbytes in files:
        if not fbytes:
            continue
        safe_name = _safe_filename(fname)
        saved_name = f"{uuid.uuid4().hex}_{safe_name}"
        path = os.path.join(folder, saved_name)
        try:
            with open(path, "wb") as fh:
                fh.write(fbytes)
        except Exception:
            continue
        mime_type, _ = mimetypes.guess_type(safe_name)
        size = 0
        try:
            size = os.path.getsize(path)
        except Exception:
            size = 0
        stored.append({
            "name": fname,
            "size": size,
            "type": mime_type or "",
            "url": f"/static/bid_documents/{bid_id}/{saved_name}",
            "uploaded_at": datetime.utcnow().isoformat(),
        })
    return stored

def _merge_uploaded_files(existing: list, incoming: list) -> list:
    merged: list = []
    seen: set = set()
    for f in (existing or []):
        if not isinstance(f, dict):
            continue
        key = (str(f.get("name") or ""), int(f.get("size") or 0))
        if key in seen:
            continue
        seen.add(key)
        merged.append(f)
    for f in (incoming or []):
        if not isinstance(f, dict):
            continue
        key = (str(f.get("name") or ""), int(f.get("size") or 0))
        if key in seen:
            continue
        seen.add(key)
        merged.append(f)
    return merged

def _attach_uploaded_files_to_bid(conn, bid_id: int, files: list[tuple[str, bytes]]) -> None:
    if not bid_id or not files:
        return
    uploaded = _store_bid_documents(bid_id, files)
    if not uploaded:
        return
    _ensure_bid_assign_meta_table(conn)
    with conn.cursor() as c:
        c.execute("SELECT data FROM bid_assign_meta WHERE g_id=%s", (int(bid_id),))
        row = c.fetchone() or {}
        try:
            meta = json.loads(row.get("data") or "{}")
        except Exception:
            meta = {}
        if not isinstance(meta, dict):
            meta = {}
        existing_files = meta.get("uploaded_files") if isinstance(meta, dict) else []
        meta["uploaded_files"] = _merge_uploaded_files(existing_files, uploaded)
        c.execute(
            """
            INSERT INTO bid_assign_meta (g_id, data)
            VALUES (%s, %s)
            ON DUPLICATE KEY UPDATE data=VALUES(data), updated_at=CURRENT_TIMESTAMP
            """,
            (int(bid_id), json.dumps(meta)),
        )
    conn.commit()

def _ensure_onedrive_tables(conn):
    with conn.cursor() as c:
        c.execute("""
        CREATE TABLE IF NOT EXISTS onedrive_folder_state (
          id INT AUTO_INCREMENT PRIMARY KEY,
          folder_path VARCHAR(1024) NOT NULL UNIQUE,
          folder_label VARCHAR(255) NULL,
          last_signature VARCHAR(64) NULL,
          last_scan_at TIMESTAMP NULL,
          last_eval_at TIMESTAMP NULL,
          created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        ) ENGINE=InnoDB CHARSET=utf8mb4;""")
    conn.commit()

def _get_onedrive_folder_state(conn, folder_path: str) -> dict | None:
    _ensure_onedrive_tables(conn)
    try:
        with conn.cursor() as c:
            c.execute(
                "SELECT * FROM onedrive_folder_state WHERE folder_path=%s LIMIT 1",
                (folder_path,),
            )
            return c.fetchone()
    except Exception:
        return None

def _upsert_onedrive_folder_state(conn, folder_path: str, *, folder_label: str | None, signature: str | None, scanned: bool, evaluated: bool) -> None:
    _ensure_onedrive_tables(conn)
    now = datetime.now()
    scan_ts = now if scanned else None
    eval_ts = now if evaluated else None
    with conn.cursor() as c:
        c.execute(
            (
                "INSERT INTO onedrive_folder_state (folder_path, folder_label, last_signature, last_scan_at, last_eval_at) "
                "VALUES (%s,%s,%s,%s,%s) "
                "ON DUPLICATE KEY UPDATE folder_label=VALUES(folder_label), last_signature=VALUES(last_signature), "
                "last_scan_at=IF(VALUES(last_scan_at) IS NULL, last_scan_at, VALUES(last_scan_at)), "
                "last_eval_at=IF(VALUES(last_eval_at) IS NULL, last_eval_at, VALUES(last_eval_at))"
            ),
            (folder_path, folder_label, signature, scan_ts, eval_ts),
        )
    conn.commit()

def _drop_all_tables(conn):
    with conn.cursor() as c:
        # Drop in safe order (no FKs defined, but keep logical order)
        for tbl in [
            "bid_incoming",
            "evaluation_runs",
            "uploaded_rfp_files",
            "company_web_context",
            "company_locations",
            "company_preferences",
            "company_capabilities",
            "company_details",
        ]:
            try:
                c.execute(f"DROP TABLE IF EXISTS {tbl}")
            except Exception:
                pass
    conn.commit()

def reset_database_schema(drop_first: bool = True):
    ok, conn, _, _ = _open_mysql_or_create()
    if not ok:
        return False, "db connect failed"
    try:
        if drop_first:
            _drop_all_tables(conn)
        _ensure_company_tables(conn)
        _ensure_scrape_tables(conn)
        _ensure_eval_tables(conn)
        return True, "ok"
    except Exception as e:
        return False, str(e)
    finally:
        try:
            conn.close()
        except Exception:
            pass

def save_full_result_to_db(
    file_name: str | None,
    result_text: str,
    export: dict | None,
    file_hash: str | None = None,
    *,
    upsert_by_file_name: bool | None = None,
) -> tuple[bool, str]:
    ok, conn, _, _ = _open_mysql_or_create()
    if not ok:
        return False, "db connect failed"
    _ensure_eval_tables(conn)
    try:
        if upsert_by_file_name is None:
            upsert_by_file_name = bool(file_name and str(file_name).startswith("FOLDER_"))
        if upsert_by_file_name and file_name:
            existing_id = None
            with conn.cursor() as c:
                c.execute(
                    "SELECT id FROM evaluation_runs WHERE file_name=%s ORDER BY id DESC LIMIT 1",
                    (file_name,),
                )
                row = c.fetchone()
                if row and row.get("id") is not None:
                    existing_id = int(row["id"])
            if existing_id is not None:
                with conn.cursor() as c:
                    c.execute(
                        "UPDATE evaluation_runs SET file_hash=%s, result_text=%s, export_json=%s, created_at=CURRENT_TIMESTAMP WHERE id=%s",
                        (file_hash or None, result_text, json.dumps(export) if export else None, existing_id),
                    )
            else:
                with conn.cursor() as c:
                    c.execute(
                        "INSERT INTO evaluation_runs (file_name, file_hash, result_text, export_json) VALUES (%s,%s,%s,%s)",
                        (file_name or None, file_hash or None, result_text, json.dumps(export) if export else None),
                    )
        else:
            with conn.cursor() as c:
                c.execute(
                    "INSERT INTO evaluation_runs (file_name, file_hash, result_text, export_json) VALUES (%s,%s,%s,%s)",
                    (file_name or None, file_hash or None, result_text, json.dumps(export) if export else None),
                )
        conn.commit()
        return True, "ok"
    except Exception as e:
        return False, str(e)


def is_due_date_past(due_date_str: str) -> bool:
    try:
        if not due_date_str or due_date_str in ("Not Found", "Not specified"):
            return False
        s = str(due_date_str).strip()[:10]
        dd = datetime.strptime(s, "%Y-%m-%d").date()
        return dd < date.today()
    except Exception:
        return False

def _ensure_go_nogo_tables(conn) -> None:
    with conn.cursor() as c:
        c.execute("""
        CREATE TABLE IF NOT EXISTS go_bids (
          g_id INT AUTO_INCREMENT PRIMARY KEY,
          bid_incoming_id INT NULL,
          id INT NULL,
          b_name VARCHAR(255),
          due_date VARCHAR(255),
          state VARCHAR(64),
          scope TEXT,
          type VARCHAR(255),
          scoring INT,
          comp_name VARCHAR(255),
          company VARCHAR(255),
          decision VARCHAR(16),
          summary LONGTEXT,
          created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
          UNIQUE KEY uniq_bname_due_date (b_name, due_date)
        ) ENGINE=InnoDB CHARSET=utf8mb4;""")
        # NOTE: Align with the phpMyAdmin schema you showed (id, bid_incoming_id, original_decision, moved_reason, moved_at).
        c.execute("""
        CREATE TABLE IF NOT EXISTS nogo_bids (
          id INT AUTO_INCREMENT PRIMARY KEY,
          bid_incoming_id INT NULL,
          b_name VARCHAR(255),
          due_date VARCHAR(255),
          state VARCHAR(64),
          scope TEXT,
          type VARCHAR(255),
          scoring INT,
          comp_name VARCHAR(255),
          original_decision VARCHAR(16),
          decision VARCHAR(16),
          summary LONGTEXT,
          moved_reason VARCHAR(64),
          moved_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
          UNIQUE KEY uniq_bid_incoming_id (bid_incoming_id),
          KEY idx_bname_due_date (b_name, due_date)
        ) ENGINE=InnoDB CHARSET=utf8mb4;""")

def _get_table_columns(conn, table: str) -> set[str]:
    try:
        with conn.cursor() as c:
            c.execute(f"SHOW COLUMNS FROM {table}")
            rows = c.fetchall() or []
        return {r.get("Field") for r in rows if r.get("Field")}
    except Exception:
        return set()

def _schema_aware_upsert(conn, table: str, data: dict) -> None:
    cols = _get_table_columns(conn, table)
    if not cols:
        return
    payload = {k: v for k, v in (data or {}).items() if k in cols}
    if not payload:
        return
    insert_cols = list(payload.keys())
    placeholders = ", ".join(["%s"] * len(insert_cols))
    insert_sql_cols = ", ".join(insert_cols)
    # Exclude only immutable/auto columns; allow moved_at to update on re-move.
    update_cols = [c for c in insert_cols if c not in ("id", "g_id", "ng_id", "created_at")]
    if update_cols:
        update_sql = ", ".join([f"{c}=VALUES({c})" for c in update_cols])
        sql = f"INSERT INTO {table} ({insert_sql_cols}) VALUES ({placeholders}) ON DUPLICATE KEY UPDATE {update_sql}"
    else:
        sql = f"INSERT INTO {table} ({insert_sql_cols}) VALUES ({placeholders})"
    with conn.cursor() as c:
        c.execute(sql, tuple(payload[k] for k in insert_cols))

def _schema_aware_delete(conn, table: str, *, bid_id: int | None, b_name: str | None, due_date: str | None) -> None:
    cols = _get_table_columns(conn, table)
    if not cols:
        return
    where = None
    params: tuple = ()
    if bid_id is not None and "bid_incoming_id" in cols:
        where = "bid_incoming_id=%s"
        params = (bid_id,)
    elif bid_id is not None and "id" in cols:
        where = "id=%s"
        params = (bid_id,)
    elif b_name and due_date and ("b_name" in cols) and ("due_date" in cols):
        where = "b_name=%s AND due_date=%s"
        params = (b_name, due_date)
    if where:
        with conn.cursor() as c:
            c.execute(f"DELETE FROM {table} WHERE {where}", params)

def _ensure_ik_project_sequence_table(conn) -> None:
    with conn.cursor() as c:
        c.execute(
            """
            CREATE TABLE IF NOT EXISTS ik_project_sequence (
              code_prefix VARCHAR(10) PRIMARY KEY,
              last_num INT NOT NULL
            ) ENGINE=InnoDB CHARSET=utf8mb4;
            """
        )
    conn.commit()

def _infer_ik_project_prefix(scope_text: str | None, type_text: str | None = None, summary_text: str | None = None) -> str | None:
    # Always consider summary as a fallback signal
    blob = f"{scope_text or ''} {type_text or ''} {summary_text or ''}".lower()
    if "solar" in blob:
        return "SLR"
    if "hvac" in blob:
        return "HVC"
    if "light" in blob or "lighting" in blob or "led" in blob:
        return "LTG"
    return None

def _next_ik_project_number(conn, prefix: str) -> int | None:
    if not prefix:
        return None
    _ensure_ik_project_sequence_table(conn)
    with conn.cursor() as c:
        c.execute(
            """
            INSERT INTO ik_project_sequence (code_prefix, last_num)
            VALUES (%s, %s)
            ON DUPLICATE KEY UPDATE last_num = last_num
            """,
            (prefix, 847000),
        )
        c.execute(
            "UPDATE ik_project_sequence SET last_num = last_num + 1 WHERE code_prefix = %s",
            (prefix,),
        )
        c.execute("SELECT last_num FROM ik_project_sequence WHERE code_prefix = %s", (prefix,))
        row = c.fetchone() or {}
    try:
        return int(row.get("last_num") or 0) or None
    except Exception:
        return None

def _ensure_ik_project_code_for_incoming(conn, bid_id: int | None, row: dict) -> str | None:
    log = _get_onedrive_logger()
    if not bid_id:
        log.warning("IK code not generated: missing bid_id")
        return None
    incoming_code = (row.get("ik_project_code") or "").strip()
    if incoming_code:
        with conn.cursor() as c:
            c.execute("UPDATE bid_incoming SET ik_project_code=%s WHERE id=%s", (incoming_code, bid_id))
        conn.commit()
        log.info("IK code already present: %s (bid_id=%s)", incoming_code, bid_id)
        return incoming_code
    prefix = _infer_ik_project_prefix(row.get("scope"), row.get("type"), row.get("summary"))
    if not prefix:
        log.warning(
            "IK code not generated: could not infer prefix. bid_id=%s scope=%s type=%s summary=%s",
            bid_id,
            (row.get("scope") or "")[:120],
            (row.get("type") or "")[:120],
            (row.get("summary") or "")[:120],
        )
        return None
    seq_num = _next_ik_project_number(conn, prefix)
    if not seq_num:
        log.error("IK code not generated: sequence unavailable. bid_id=%s prefix=%s", bid_id, prefix)
        return None
    project_code = f"IK-P/{prefix}/{seq_num:06d}"
    with conn.cursor() as c:
        c.execute("UPDATE bid_incoming SET ik_project_code=%s WHERE id=%s", (project_code, bid_id))
    conn.commit()
    log.info("IK code generated: %s (bid_id=%s)", project_code, bid_id)
    return project_code

def save_bid_result_to_db(row: dict) -> tuple[bool, str, int | None]:
    log = _get_onedrive_logger()
    ok, conn, _, _ = _open_mysql_or_create()
    if not ok:
        log.error("save_bid_result_to_db: db connect failed")
        return False, "db connect failed", None
    _ensure_eval_tables(conn)
    # Coerce payload into DB-friendly types (e.g., scoring must be INT)
    def _coerce(r: dict) -> dict:
        out = dict(r or {})
        scoring = out.get("scoring")
        if isinstance(scoring, dict):
            try:
                # use highest numeric score from dict
                vals = [v for v in scoring.values() if isinstance(v, (int, float))]
                out["scoring"] = int(round(max(vals))) if vals else None
            except Exception:
                out["scoring"] = None
        elif isinstance(scoring, (int, float)):
            out["scoring"] = int(round(scoring))
        else:
            out["scoring"] = None
        # Ensure strings
        for k in ["b_name","due_date","state","scope","type","comp_name","decision","summary","ik_project_code"]:
            if k in out and out[k] is not None and not isinstance(out[k], str):
                out[k] = str(out[k])
        return out
    row = _coerce(row)
    # Normalize due_date to YYYY-MM-DD if possible
    def _normalize_due_date(d):
        try:
            if d is None:
                return None
            s = str(d).strip()
            # ISO-like: YYYY-MM-DD or YYYY-MM-DDTHH:MM:SS
            if len(s) >= 10 and s[4] == "-" and s[7] == "-":
                return s[:10]
            from datetime import datetime as dt
            for fmt in ("%B %d, %Y", "%Y-%m-%d", "%Y/%m/%d", "%m/%d/%Y", "%d/%m/%Y", "%Y-%m-%d %H:%M:%S"):
                try:
                    return dt.strptime(s, fmt).strftime("%Y-%m-%d")
                except Exception:
                    pass
            return s
        except Exception:
            return d
    row["due_date"] = _normalize_due_date(row.get("due_date"))
    evaluation_date = date.today()
    try:
        b_name = (row.get("b_name") or "").strip()
        due_date_norm = (row.get("due_date") or "").strip()
        due_date_past = is_due_date_past(due_date_norm)

        # If due date already passed, always route to NO-GO
        original_decision = (row.get("decision") or "").strip()
        effective_decision = original_decision
        if due_date_past:
            effective_decision = "NO-GO"
            row["decision"] = effective_decision

        inserted_id: int | None = None
        updated_existing = False
        existing_id: int | None = None

        # Bid-level duplicate detection: treat same (b_name, due_date) as duplicate and update instead of inserting
        if b_name and due_date_norm:
            with conn.cursor() as c:
                c.execute(
                    "SELECT id FROM bid_incoming WHERE b_name=%s AND due_date=%s ORDER BY id DESC LIMIT 1",
                    (b_name, due_date_norm),
                )
                existing = c.fetchone()
                if existing and existing.get("id") is not None:
                    existing_id = int(existing["id"])

        if existing_id is not None:
            with conn.cursor() as c:
                c.execute(
                    (
                        "UPDATE bid_incoming SET state=%s, scope=%s, type=%s, scoring=%s, comp_name=%s, ik_project_code=%s, decision=%s, summary=%s, evaluation_date=%s "
                        "WHERE id=%s"
                    ),
                    (
                        row.get("state"),
                        row.get("scope"),
                        row.get("type"),
                        row.get("scoring"),
                        row.get("comp_name"),
                        row.get("ik_project_code"),
                        effective_decision,
                        row.get("summary"),
                        evaluation_date,
                        existing_id,
                    ),
                )
            updated_existing = True
        else:
            with conn.cursor() as c:
                c.execute(
                    (
                        "INSERT INTO bid_incoming (b_name, due_date, state, scope, type, scoring, comp_name, ik_project_code, decision, summary, evaluation_date) "
                        "VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
                    ),
                    (
                        row.get("b_name"),
                        row.get("due_date"),
                        row.get("state"),
                        row.get("scope"),
                        row.get("type"),
                        row.get("scoring"),
                        row.get("comp_name"),
                        row.get("ik_project_code"),
                        effective_decision,
                        row.get("summary"),
                        evaluation_date,
                    ),
                )
                inserted_id = int(c.lastrowid)

        conn.commit()
        bid_id = existing_id if existing_id is not None else inserted_id
        try:
            _ensure_ik_project_code_for_incoming(conn, bid_id, row)
        except Exception:
            log.exception("IK code generation failed: bid_id=%s", bid_id)
            pass

        # Mirror GO/NO-GO decisions into go_bids / nogo_bids (schema-aware)
        try:
            _ensure_go_nogo_tables(conn)
            base_payload = {
                "bid_incoming_id": bid_id,
                "id": bid_id,
                "b_name": row.get("b_name"),
                "due_date": row.get("due_date"),
                "state": row.get("state"),
                "scope": row.get("scope"),
                "type": row.get("type"),
                "scoring": row.get("scoring"),
                "comp_name": row.get("comp_name"),
                "company": row.get("comp_name"),
                "ik_project_code": row.get("ik_project_code"),
                "summary": row.get("summary"),
            }

            decision_lc = (effective_decision or "").strip().lower()
            if decision_lc == "go" and not due_date_past:
                _schema_aware_upsert(conn, "go_bids", {**base_payload, "decision": effective_decision})
                _schema_aware_delete(conn, "nogo_bids", bid_id=bid_id, b_name=b_name, due_date=due_date_norm)
            elif decision_lc in ("no-go", "nogo", "no go", "no_go", "nogobid") or due_date_past:
                moved_reason = "due_date_passed" if due_date_past else "no_go_decision"
                _schema_aware_upsert(
                    conn,
                    "nogo_bids",
                    {
                        **base_payload,
                        "decision": effective_decision,
                        "original_decision": original_decision,
                        "moved_reason": moved_reason,
                        "moved_at": datetime.now(),
                    },
                )
                _schema_aware_delete(conn, "go_bids", bid_id=bid_id, b_name=b_name, due_date=due_date_norm)
            conn.commit()
        except Exception as e:
            logging.exception("Failed mirroring bid into go/nogo tables: %s", e)

        if updated_existing:
            log.info("save_bid_result_to_db: updated existing bid_incoming id=%s b_name=%s due_date=%s", bid_id, b_name, due_date_norm)
            return True, "duplicate detected: updated existing bid_incoming row", bid_id
        log.info("save_bid_result_to_db: inserted bid_incoming id=%s b_name=%s due_date=%s", bid_id, b_name, due_date_norm)
        return True, "ok", bid_id
    except Exception as e:
        log.exception("save_bid_result_to_db failed")
        return False, str(e), None

def parse_db_export_line(text: str) -> dict | None:
    try:
        for line in reversed((text or "").splitlines()):
            if line.strip().startswith("DB_EXPORT_JSON:"):
                payload = line.split(":", 1)[1].strip()
                return json.loads(payload)
    except Exception:
        return None
    return None

# Attempts to find the first valid top-level JSON object anywhere in text and parse it
def parse_first_json_object(text: str) -> dict | None:
    s = text or ""
    if not s:
        return None
    n = len(s)
    i = s.find("{")
    while i != -1 and i < n:
        depth = 0
        in_str = False
        escape = False
        j = i
        while j < n:
            ch = s[j]
            if in_str:
                if escape:
                    escape = False
                elif ch == "\\":
                    escape = True
                elif ch == '"':
                    in_str = False
            else:
                if ch == '"':
                    in_str = True
                elif ch == '{':
                    depth += 1
                elif ch == '}':
                    depth -= 1
                    if depth == 0:
                        candidate = s[i:j+1]
                        try:
                            return json.loads(candidate)
                        except Exception:
                            break
            j += 1
        i = s.find("{", i + 1)
    return None

# ----------------- REPORT SECTION EXTRACTORS -----------------
def extract_final_recommendation_summary(text: str) -> str | None:
    """Extract the FINAL RECOMMENDATION narrative from the rendered report."""
    if not text:
        return None
    try:
        pattern = re.compile(
            r"(?is)"
            r"(?:^|\n)#{0,3}\s*final\s+recommendation[^\n]*\n+"
            r"(.+?)"
            r"(?=\n\s*(?:#{1,3}\s*[A-Z].+|database\s+export\s+json|db_export_json\s*:)|\Z)"
        )
        m = pattern.search(text)
        if m:
            body = m.group(1).strip()
            body = re.sub(r"[ \t]+\n", "\n", body)
            body = re.sub(r"\n{3,}", "\n\n", body)
            return body[:4000]
        # Fallback to line with Executive Takeaway if present
        m2 = re.search(r"(?is)(Executive\s+Takeaway:.*)", text)
        if m2:
            return m2.group(1).strip()[:4000]
        return None
    except Exception:
        return None

def extract_scope_of_work_section(text: str) -> str | None:
    """Extract the SCOPE OF WORK section that follows the Final Recommendation."""
    if not text:
        return None
    try:
        pattern = re.compile(
            r"(?is)"
            r"(?:^|\n)#{0,3}\s*scope\s+of\s+work[^\n]*\n+"
            r"(.+?)"
            r"(?=\n\s*(?:#{1,3}\s*[A-Z].+|database\s+export\s+json|db_export_json\s*:)|\Z)"
        )
        m = pattern.search(text)
        if m:
            body = m.group(1).strip()
            body = re.sub(r"[ \t]+\n", "\n", body)
            body = re.sub(r"\n{3,}", "\n\n", body)
            return body[:4000]
        return None
    except Exception:
        return None

def extract_compliance_table_section(text: str) -> str | None:
    """Extract the COMPLIANCE TABLE section (often used as 'context' before the final recommendation)."""
    if not text:
        return None
    try:
        pattern = re.compile(
            r"(?is)"
            r"(?:^|\n)#{0,3}\s*compliance\s+table[^\n]*\n+"
            r"(.+?)"
            r"(?=\n\s*(?:#{1,3}\s*[A-Z].+|final\s+recommendation|scope\s+of\s+work|database\s+export\s+json|db_export_json\s*:)|\Z)"
        )
        m = pattern.search(text)
        if m:
            body = m.group(1).strip()
            body = re.sub(r"[ \t]+\n", "\n", body)
            body = re.sub(r"\n{3,}", "\n\n", body)
            return body[:4000]
        return None
    except Exception:
        return None

def extract_step1_data_extraction_section(text: str) -> str | None:
    """Extract STEP 1 - RFP DATA EXTRACTION section (table)."""
    if not text:
        return None
    try:
        pattern = re.compile(
            r"(?is)"
            r"(?:^|\n)#{0,3}\s*step\s*1[^\n]*rfp\s*data\s*extraction[^\n]*\n+"
            r"(.+?)"
            r"(?=\n\s*(?:#{1,3}\s*[A-Z].+|company\s+selection\s+criteria|detailed\s+bid\s+evaluation|compliance\s+table|final\s+recommendation|scope\s+of\s+work|database\s+export\s+json|db_export_json\s*:)|\Z)"
        )
        m = pattern.search(text)
        if m:
            body = m.group(1).strip()
            body = re.sub(r"[ \t]+\n", "\n", body)
            body = re.sub(r"\n{3,}", "\n\n", body)
            return body[:12000]
        return None
    except Exception:
        return None

def extract_detailed_bid_evaluation_section(text: str) -> str | None:
    """Extract DETAILED BID EVALUATION section (all companies, a..i scoring/remarks)."""
    if not text:
        return None
    try:
        pattern = re.compile(
            r"(?is)"
            r"(?:^|\n)#{0,3}\s*detailed\s+bid\s+evaluation[^\n]*\n+"
            r"(.+?)"
            r"(?=\n\s*(?:#{1,3}\s*[A-Z].+|compliance\s+table|final\s+recommendation|scope\s+of\s+work|database\s+export\s+json|db_export_json\s*:)|\Z)"
        )
        m = pattern.search(text)
        if m:
            body = m.group(1).strip()
            body = re.sub(r"[ \t]+\n", "\n", body)
            body = re.sub(r"\n{3,}", "\n\n", body)
            return body[:20000]
        return None
    except Exception:
        return None

def extract_evaluation_sequence_section(text: str) -> str | None:
    """Extract EVALUATION SEQUENCE section (Supply/Installation/Substitution)."""
    if not text:
        return None
    try:
        pattern = re.compile(
            r"(?is)"
            r"(?:^|\n)#{0,3}\s*evaluation\s+sequence[^\n]*\n+"
            r"(.+?)"
            r"(?=\n\s*(?:#{1,3}\s*[A-Z].+|detailed\s+bid\s+evaluation|compliance\s+table|final\s+recommendation|scope\s+of\s+work|database\s+export\s+json|db_export_json\s*:)|\Z)"
        )
        m = pattern.search(text)
        if m:
            body = m.group(1).strip()
            body = re.sub(r"[ \t]+\n", "\n", body)
            body = re.sub(r"\n{3,}", "\n\n", body)
            return body[:8000]
        return None
    except Exception:
        return None

def build_scope_of_work_fallback(rfp_text: str, profiles: list[str]) -> str:
    """Compose a concise scope of work from detected profiles and requirement flags as fallback."""
    t = (rfp_text or "").lower()
    criteria = analyze_work_criteria(rfp_text)
    parts: list[str] = []
    if profiles:
        parts.append(f"Work Type: {', '.join(profiles)}.")
    delivery_bits: list[str] = []
    if criteria.get("supply"):
        delivery_bits.append("Supply")
    if criteria.get("installation"):
        delivery_bits.append("Installation")
    subs = criteria.get("substitution")
    if subs and subs != "Not specified":
        delivery_bits.append(f"Substitution: {subs}")
    if delivery_bits:
        parts.append("Delivery: " + ", ".join(delivery_bits) + ".")
    # Requirement flags
    site = _has_any(t, ["site visit", "site investigation", "pre-bid meeting", "walkthrough"])
    baba = _has_any(t, ["baba", "build america", "baa", "buy american", "davis bacon"])
    bonds = _has_any(t, ["bid bond", "performance bond", "payment bond"])
    ins = _has_any(t, ["insurance", "general liability", "workers compensation"])
    req_bits: list[str] = []
    if site: req_bits.append("Pre-bid site visit/investigation")
    if baba: req_bits.append("BABA/BAA/Davis-Bacon compliance")
    if bonds: req_bits.append("Bid/Performance/Payment bonds")
    if ins: req_bits.append("Insurance and COI")
    if req_bits:
        parts.append("Requirements: " + ", ".join(req_bits) + ".")
    return " ".join(parts).strip()

# ----------------- OCR EXTRACTION -----------------
@st.cache_resource
def _get_doctr_predictor():
    if not _DOCTR_OK:
        return None
    try:
        return ocr_predictor(pretrained=True)
    except Exception:
        return None

@st.cache_data(show_spinner=False)
def extract_text(file_bytes: bytes, filename: str) -> str:
    name = (filename or "").lower()
    # TXT fast path
    if name.endswith(".txt"):
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            return ""
    # DOCX native path if available
    if name.endswith(".docx") and DocxDocument is not None:
        try:
            doc = DocxDocument(BytesIO(file_bytes))
            return "\n".join([(p.text or "").strip() for p in doc.paragraphs if (p.text or "").strip()])
        except Exception:
            pass
    # OCR path
    if _DOCTR_OK:
        predictor = _get_doctr_predictor()
        if predictor is not None:
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            try:
                if name.endswith(".pdf"):
                    doc = DocumentFile.from_pdf(tmp_path)
                else:
                    # Pass file path directly to avoid unsupported object types
                    doc = DocumentFile.from_images([tmp_path])
                result = predictor(doc)
                export = result.export()
                text = "\n".join(
                    " ".join(w.get("value", "") for w in line.get("words", []))
                    for page in export.get("pages", [])
                    for block in page.get("blocks", [])
                    for line in block.get("lines", [])
                )
                return text
            finally:
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass
    # Fallback best-effort
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        return ""

# ----------------- COMPANY DB CONTEXT -----------------
def _is_safe_sql_ident(name: str) -> bool:
    """Best-effort validation for table/column identifiers before interpolating into SQL."""
    return bool(re.match(r"^[A-Za-z0-9_]+$", str(name or "")))


def _discover_case_studies_table(conn) -> tuple[str | None, str | None, str | None, list[str]]:
    """
    Try to find a 'case studies' table and identify:
    - table name
    - company id column (preferred)
    - company name column (fallback)
    - text-like columns to extract into context/scoring
    """
    candidates: list[str] = []
    try:
        with conn.cursor() as c:
            # Prefer explicit names first
            for tname in ["case_studies", "company_case_studies", "company_casestudies", "company_case_study", "case_study"]:
                try:
                    c.execute("SHOW TABLES LIKE %s", (tname,))
                    if c.fetchone():
                        candidates.append(tname)
                except Exception:
                    pass
            # Then try discovery via information_schema
            if not candidates:
                c.execute(
                    """
                    SELECT table_name
                    FROM information_schema.tables
                    WHERE table_schema = DATABASE()
                      AND table_name LIKE %s
                      AND table_name LIKE %s
                    """,
                    ("%case%", "%stud%"),
                )
                for r in (c.fetchall() or []):
                    tn = r.get("table_name")
                    if tn and tn not in candidates:
                        candidates.append(tn)
    except Exception:
        return None, None, None, []

    for table in candidates:
        if not _is_safe_sql_ident(table):
            continue
        try:
            with conn.cursor() as c:
                c.execute(f"SHOW COLUMNS FROM {table}")
                col_rows = (c.fetchall() or [])
                cols = [r.get("Field") for r in col_rows if r.get("Field")]
        except Exception:
            # SHOW COLUMNS failed or table unreadable
            continue

        cols_l = [str(x).lower() for x in cols]
        id_col = None
        name_col = None
        # Prefer company_id style
        for cand in ["company_id", "companyid", "comp_id", "companydetails_id", "company_details_id"]:
            if cand in cols_l:
                id_col = cols[cols_l.index(cand)]
                break
        # Fallback name column
        for cand in ["company_name", "company", "name", "vendor_name"]:
            if cand in cols_l:
                name_col = cols[cols_l.index(cand)]
                break

        # Choose text-like columns (exclude ids and company refs)
        skip = {"id", "created_at", "updated_at", "last_updated", "timestamp"}
        if id_col:
            skip.add(str(id_col).lower())
        if name_col:
            skip.add(str(name_col).lower())
        text_cols: list[str] = []
        try:
            with conn.cursor() as c:
                c.execute(f"SHOW COLUMNS FROM {table}")
                for r in (c.fetchall() or []):
                    field = r.get("Field")
                    f_l = (field or "").lower()
                    if not field or f_l in skip:
                        continue
                    typ = (r.get("Type") or "").lower()
                    if any(x in typ for x in ["text", "char", "varchar", "longtext", "mediumtext", "tinytext"]):
                        text_cols.append(field)
        except Exception:
            text_cols = []

        return table, id_col, name_col, text_cols

    return None, None, None, []


def _load_case_studies_into_sections(conn, sections_by_name: dict) -> None:
    """Mutates sections_by_name[...] to attach 'case_studies' list if a case studies table exists."""
    table, id_col, name_col, _text_cols = _discover_case_studies_table(conn)
    if not table:
        return
    if not _is_safe_sql_ident(table):
        return
    if id_col and not _is_safe_sql_ident(id_col):
        id_col = None
    if name_col and not _is_safe_sql_ident(name_col):
        name_col = None

    # Ensure key exists
    for cname in sections_by_name:
        sections_by_name[cname].setdefault("case_studies", [])

    # Build lookup sets
    name_set_lower = {str(n or "").strip().lower() for n in sections_by_name.keys()}
    id_to_name: dict[int, str] = {}
    for cname, sec in sections_by_name.items():
        cid = sec.get("company_id")
        if isinstance(cid, int):
            id_to_name[cid] = cname

    try:
        with conn.cursor() as c:
            if id_col and id_to_name:
                ids = list(id_to_name.keys())
                placeholders = ",".join(["%s"] * len(ids))
                c.execute(f"SELECT * FROM {table} WHERE {id_col} IN ({placeholders})", tuple(ids))
                for r in (c.fetchall() or []):
                    cid = r.get(id_col)
                    cname = id_to_name.get(cid) if isinstance(cid, int) else None
                    if cname and cname in sections_by_name:
                        sections_by_name[cname]["case_studies"].append(r)
                return

            # Fallback: match by name if available
            if name_col and name_set_lower:
                # Pull only relevant rows by filtering in python (case-insensitive), to avoid LOWER() preventing indexes
                c.execute(f"SELECT * FROM {table}")
                for r in (c.fetchall() or []):
                    nm = str(r.get(name_col) or "").strip().lower()
                    if nm in name_set_lower:
                        # Find original casing key
                        for cname in sections_by_name:
                            if cname.strip().lower() == nm:
                                sections_by_name[cname]["case_studies"].append(r)
                                break
    except Exception:
        return


def _discover_tables_with_column(conn, col_names_lower: set[str]) -> list[str]:
    """Return table names in current schema that contain any column in col_names_lower."""
    out: list[str] = []
    try:
        with conn.cursor() as c:
            placeholders = ",".join(["%s"] * len(col_names_lower))
            c.execute(
                f"""
                SELECT DISTINCT table_name
                FROM information_schema.columns
                WHERE table_schema = DATABASE()
                  AND LOWER(column_name) IN ({placeholders})
                """,
                tuple(col_names_lower),
            )
            for r in (c.fetchall() or []):
                tn = r.get("table_name")
                if tn and tn not in out:
                    out.append(tn)
    except Exception:
        return []
    return out


def _should_include_context_table(table_name: str) -> bool:
    """
    Only include likely company-context tables.
    This avoids pulling transactional/bid tables that can be huge or irrelevant.
    """
    t = (table_name or "").lower().strip()
    if not t:
        return False
    # hard deny list / patterns
    deny_prefixes = (
        "bid_", "bids", "go_", "nogo_", "evaluation_", "uploaded_", "logs", "action_", "assigned_",
        "departments", "employees", "permissions", "triggers", "tracking",
    )
    if t.startswith(deny_prefixes):
        return False
    # allow common context patterns
    allow_prefixes = (
        "company_", "case_", "past_", "performance", "capabil", "prefer", "location", "web_",
    )
    return t.startswith(allow_prefixes) or ("case" in t and "study" in t) or ("past" in t and "perform" in t)


def _select_reasonable_columns(conn, table: str) -> tuple[list[str], dict]:
    """
    Choose columns that are useful for context:
    - keep ids/foreign keys (company_id/case_study_id) for routing
    - keep text/varchar
    - keep numeric/date fields (they help with 'past performance' evidence)
    """
    cols: list[str] = []
    meta: dict = {}
    try:
        with conn.cursor() as c:
            c.execute(f"SHOW COLUMNS FROM {table}")
            rows = (c.fetchall() or [])
        for r in rows:
            field = r.get("Field")
            typ = (r.get("Type") or "").lower()
            if not field:
                continue
            f_l = str(field).lower()
            # always keep routing columns and primary key-ish ids
            if f_l in {"id", "company_id", "companyid", "company_name", "case_study_id", "casestudy_id"}:
                cols.append(field)
                continue
            # skip obvious noise
            if f_l in {"created_at", "updated_at", "last_fetched", "password", "token"}:
                continue
            # keep text-like and numeric/date-like
            if any(x in typ for x in ["text", "char", "varchar", "longtext", "mediumtext", "tinytext"]):
                cols.append(field)
            elif any(x in typ for x in ["int", "decimal", "float", "double", "date", "time", "year"]):
                cols.append(field)
        # cap to avoid giant SELECT lists
        if len(cols) > 40:
            cols = cols[:40]
        meta = {"types": {r.get("Field"): (r.get("Type") or "") for r in rows}}
    except Exception:
        return [], {}
    return cols, meta


def _row_to_compact_kv(row: dict, drop_keys_lower: set[str], max_field_chars: int = 800) -> str:
    items: list[str] = []
    for k, v in (row or {}).items():
        kl = str(k).lower()
        if kl in drop_keys_lower:
            continue
        if v is None:
            continue
        s = str(v).strip()
        if not s:
            continue
        if len(s) > max_field_chars:
            s = s[:max_field_chars] + "…"
        items.append(f"{k}: {s}")
    return " | ".join(items)


def _load_company_linked_tables_into_sections(conn, sections_by_name: dict) -> None:
    """
    Attach extra DB tables into sections to improve scoring/remarks.
    - Loads tables with company_id/company_name columns
    - Also loads case_study_* child tables by case_study_id when present
    """
    # ensure containers
    for cname in sections_by_name:
        sections_by_name[cname].setdefault("extra_tables", {})  # {table: [rows]}

    # build lookups
    id_to_name: dict[int, str] = {}
    name_to_name: dict[str, str] = {}
    for cname, sec in sections_by_name.items():
        if isinstance(sec.get("company_id"), int):
            id_to_name[int(sec["company_id"])] = cname
        name_to_name[str(cname).strip().lower()] = cname

    company_id_tables = _discover_tables_with_column(conn, {"company_id", "companyid", "comp_id"})
    company_name_tables = _discover_tables_with_column(conn, {"company_name", "company", "vendor_name", "name"})
    candidate_tables = list(dict.fromkeys(company_id_tables + company_name_tables))  # stable dedupe

    # Don't re-load tables already handled explicitly
    already_handled = {
        "company_details", "company_capabilities", "company_preferences", "company_locations", "company_web_context",
    }

    for table in candidate_tables:
        if not table or table in already_handled:
            continue
        if not _is_safe_sql_ident(table) or not _should_include_context_table(table):
            continue

        # Identify routing columns
        try:
            with conn.cursor() as c:
                c.execute(f"SHOW COLUMNS FROM {table}")
                cols = [r.get("Field") for r in (c.fetchall() or []) if r.get("Field")]
        except Exception:
            continue
        cols_l = [str(x).lower() for x in cols]
        company_id_col = None
        company_name_col = None
        case_study_id_col = None
        for cand in ["company_id", "companyid", "comp_id"]:
            if cand in cols_l:
                company_id_col = cols[cols_l.index(cand)]
                break
        for cand in ["company_name", "company", "vendor_name", "name"]:
            if cand in cols_l:
                company_name_col = cols[cols_l.index(cand)]
                break
        for cand in ["case_study_id", "casestudy_id"]:
            if cand in cols_l:
                case_study_id_col = cols[cols_l.index(cand)]
                break

        # Select columns to pull (smaller + more relevant)
        sel_cols, _meta = _select_reasonable_columns(conn, table)
        if not sel_cols:
            continue
        sel_cols_safe = [c for c in sel_cols if _is_safe_sql_ident(c)]
        if not sel_cols_safe:
            continue
        sel_sql = ", ".join(sel_cols_safe)

        # Query by company_id in one go if possible
        try:
            with conn.cursor() as c:
                rows = []
                if company_id_col and id_to_name:
                    ids = list(id_to_name.keys())
                    placeholders = ",".join(["%s"] * len(ids))
                    order_clause = "ORDER BY id DESC" if "id" in cols_l else ""
                    c.execute(
                        f"SELECT {sel_sql} FROM {table} WHERE {company_id_col} IN ({placeholders}) {order_clause} LIMIT 300",
                        tuple(ids),
                    )
                    rows = (c.fetchall() or [])
                    for r in rows:
                        cid = r.get(company_id_col)
                        cname = id_to_name.get(cid) if isinstance(cid, int) else None
                        if cname:
                            sections_by_name[cname]["extra_tables"].setdefault(table, []).append(r)
                    continue

                # Fallback: match by company name (less ideal)
                if company_name_col and name_to_name:
                    # Pull limited rows and route in python (case-insensitive)
                    order_clause = "ORDER BY id DESC" if "id" in cols_l else ""
                    c.execute(f"SELECT {sel_sql} FROM {table} {order_clause} LIMIT 500")
                    rows = (c.fetchall() or [])
                    for r in rows:
                        nm = str(r.get(company_name_col) or "").strip().lower()
                        cname = name_to_name.get(nm)
                        if cname:
                            sections_by_name[cname]["extra_tables"].setdefault(table, []).append(r)
        except Exception:
            continue

    # Case-study child tables (case_study_id -> case_studies.id)
    # We already loaded case_studies rows; now try to load any tables containing case_study_id.
    cs_id_tables = _discover_tables_with_column(conn, {"case_study_id", "casestudy_id"})
    cs_id_tables = [t for t in cs_id_tables if t and _is_safe_sql_ident(t) and _should_include_context_table(t)]

    # Build case_study_id -> company mapping once
    csid_to_company: dict[int, str] = {}
    for cname, sec in sections_by_name.items():
        for cs in (sec.get("case_studies") or []):
            if isinstance(cs, dict) and isinstance(cs.get("id"), int):
                csid_to_company[int(cs["id"])] = cname
    if not csid_to_company:
        return

    for table in cs_id_tables:
        if table in {"case_studies"}:
            continue
        try:
            with conn.cursor() as c:
                c.execute(f"SHOW COLUMNS FROM {table}")
                cols = [r.get("Field") for r in (c.fetchall() or []) if r.get("Field")]
        except Exception:
            continue
        cols_l = [str(x).lower() for x in cols]
        csid_col = None
        for cand in ["case_study_id", "casestudy_id"]:
            if cand in cols_l:
                csid_col = cols[cols_l.index(cand)]
                break
        if not csid_col:
            continue
        sel_cols, _meta = _select_reasonable_columns(conn, table)
        sel_cols_safe = [c for c in (sel_cols or []) if _is_safe_sql_ident(c)]
        if not sel_cols_safe:
            continue
        sel_sql = ", ".join(sel_cols_safe)
        try:
            with conn.cursor() as c:
                ids = list(csid_to_company.keys())
                placeholders = ",".join(["%s"] * len(ids))
                order_clause = "ORDER BY id DESC" if "id" in cols_l else ""
                c.execute(
                    f"SELECT {sel_sql} FROM {table} WHERE {csid_col} IN ({placeholders}) {order_clause} LIMIT 500",
                    tuple(ids),
                )
                rows = (c.fetchall() or [])
                for r in rows:
                    csid = r.get(csid_col)
                    cname = csid_to_company.get(csid) if isinstance(csid, int) else None
                    if cname:
                        sections_by_name[cname]["extra_tables"].setdefault(table, []).append(r)
        except Exception:
            continue


def load_company_sections_from_db(company_names: List[str]) -> Dict[str, Dict]:
    ok, conn, _, _ = _open_mysql_or_create()
    if not ok:
        return {}
    _ensure_company_tables(conn)
    _ensure_scrape_tables(conn)
    data: Dict[str, Dict] = {
        cname: {
            "company_id": None,
            "details": None,
            "capabilities": [],
            "preferences": None,
            "location": None,
            "web_context": None,
            "case_studies": [],
        }
        for cname in company_names
    }
    with conn.cursor() as c:
        c.execute("SELECT * FROM company_details")
        for r in c.fetchall():
            n = r.get("company_name")
            if n in data:
                data[n]["details"] = r
                if isinstance(r.get("id"), int):
                    data[n]["company_id"] = r.get("id")
        c.execute("SELECT * FROM company_capabilities")
        for r in c.fetchall():
            n = r.get("company_name")
            if n in data:
                data[n]["capabilities"].append(r)
        c.execute("SELECT * FROM company_preferences")
        for r in c.fetchall():
            n = r.get("company_name")
            if n in data:
                data[n]["preferences"] = r
        c.execute("SELECT * FROM company_locations")
        for r in c.fetchall():
            n = r.get("company_name")
            if n in data:
                data[n]["location"] = r
        c.execute("SELECT company_name, content FROM company_web_context")
        for r in c.fetchall():
            n = r.get("company_name")
            if n in data:
                data[n]["web_context"] = r.get("content")
        # Attach case studies if a case-studies table exists (keyed by company_id when possible)
        _load_case_studies_into_sections(conn, data)
        # Attach *all other* relevant company-linked tables (incl. past performance, case study child tables, etc.)
        _load_company_linked_tables_into_sections(conn, data)
    conn.close()
    return data

def format_company_db_context(db_ctx: Dict[str, Dict]) -> str:
    lines = []
    for cname, sections in db_ctx.items():
        lines.append(f"COMPANY CONTEXT - {cname}")
        d = sections.get("details") or {}
        if d:
            lines.append(f"- Address: {d.get('address')}")
            lines.append(f"- Start Date: {d.get('start_date')}")
        p = sections.get("preferences") or {}
        if p:
            lines.append(f"- Deal Breakers: {p.get('deal_breakers')}")
            lines.append(f"- Deal Makers: {p.get('deal_makers')}")
            lines.append(f"- States: {p.get('preferred_states')}")
        caps = sections.get("capabilities") or []
        for c in caps:
            lines.append(f"- {c.get('capability_title')}: {c.get('capability_description')}")
    return "\n".join(lines)

# ----------------- COMPANY LOCATION HELPERS -----------------
def _get_company_base_location_from_db(company_name: str) -> dict | None:
    ok, conn, _, _ = _open_mysql_or_create()
    if not ok:
        return None
    try:
        _ensure_company_tables(conn)
        with conn.cursor() as c:
            c.execute("SELECT base_location, base_state FROM company_locations WHERE company_name=%s", (company_name,))
            row = c.fetchone()
            return row
    except Exception:
        return None
    finally:
        try:
            conn.close()
        except Exception:
            pass

def get_company_base_location(company_name: str) -> str | None:
    row = _get_company_base_location_from_db(company_name)
    if row and (row.get("base_location") or row.get("base_state")):
        # Prefer explicit base_location; fall back to state-only if provided
        if row.get("base_location"):
            return str(row.get("base_location"))
        if row.get("base_state"):
            return str(row.get("base_state"))
    # Fallback to code constant
    return COMPANY_LOCATIONS.get(company_name)

def get_company_base_state(company_name: str) -> str:
    row = _get_company_base_location_from_db(company_name)
    if row and row.get("base_state"):
        return str(row.get("base_state")).strip().upper()
    # Derive from base_location string or fallback dict
    loc = (row.get("base_location") if row else None) or COMPANY_LOCATIONS.get(company_name, ", ")
    try:
        return loc.split(",")[-1].strip().upper()
    except Exception:
        return ""

# ----------------- STATE & PROFILE -----------------
US_STATE_ABBRS = {"AL","AK","AZ","AR","CA","CO","CT","DE","FL","GA","HI","ID","IL","IN","IA","KS","KY","LA","ME","MD","MA","MI","MN",
                  "MS","MO","MT","NE","NV","NH","NJ","NM","NY","NC","ND","OH","OK","OR","PA","RI","SC","SD","TN","TX","UT","VT","VA",
                  "WA","WV","WI","WY"}

# State name to abbreviation mapping
STATE_NAME_TO_ABBR = {
    "ALABAMA": "AL", "ALASKA": "AK", "ARIZONA": "AZ", "ARKANSAS": "AR", "CALIFORNIA": "CA",
    "COLORADO": "CO", "CONNECTICUT": "CT", "DELAWARE": "DE", "FLORIDA": "FL", "GEORGIA": "GA",
    "HAWAII": "HI", "IDAHO": "ID", "ILLINOIS": "IL", "INDIANA": "IN", "IOWA": "IA",
    "KANSAS": "KS", "KENTUCKY": "KY", "LOUISIANA": "LA", "MAINE": "ME", "MARYLAND": "MD",
    "MASSACHUSETTS": "MA", "MICHIGAN": "MI", "MINNESOTA": "MN", "MISSISSIPPI": "MS", "MISSOURI": "MO",
    "MONTANA": "MT", "NEBRASKA": "NE", "NEVADA": "NV", "NEW HAMPSHIRE": "NH", "NEW JERSEY": "NJ",
    "NEW MEXICO": "NM", "NEW YORK": "NY", "NORTH CAROLINA": "NC", "NORTH DAKOTA": "ND", "OHIO": "OH",
    "OKLAHOMA": "OK", "OREGON": "OR", "PENNSYLVANIA": "PA", "RHODE ISLAND": "RI", "SOUTH CAROLINA": "SC",
    "SOUTH DAKOTA": "SD", "TENNESSEE": "TN", "TEXAS": "TX", "UTAH": "UT", "VERMONT": "VT",
    "VIRGINIA": "VA", "WASHINGTON": "WA", "WEST VIRGINIA": "WV", "WISCONSIN": "WI", "WYOMING": "WY"
}

def normalize_state_for_comparison(state: str | None) -> str:
    """Convert state name/abbreviation to 2-letter abbreviation for comparison."""
    if not state:
        return ""
    s = str(state).strip().upper()
    if len(s) == 2 and s in US_STATE_ABBRS:
        return s
    return STATE_NAME_TO_ABBR.get(s, "")

def extract_project_state_simple(text: str) -> str | None:
    if not text:
        return None
    txt = (text or "").upper()
    states_alt = "|".join(sorted(list(US_STATE_ABBRS)))
    m = re.search(r",\s*(" + states_alt + r")\b", txt)
    if m:
        return m.group(1)
    if re.search(r"\bUTAH\b", txt): return "UT"
    if re.search(r"\bTEXAS\b", txt): return "TX"
    if re.search(r"\bINDIANA\b", txt): return "IN"
    m = re.search(r"\bSTATE\s*[:\-]\s*(" + states_alt + r")\b", txt)
    if m:
        return m.group(1)
    return None

def detect_work_profiles(rfp_text: str) -> list[str]:
    t = (rfp_text or "").lower()
    detected: list[str] = []
    if "lighting" in t:
        supply = any(k in t for k in ["supply","furnish","provide materials","materials","equipment","procure","purchase","f&i","furnish and install"])
        install = any(k in t for k in ["install","installation","replace","retrofit","mount","erect"])
        no_sub = any(p in t for p in ["no substitution","substitution not allowed","no alternates","no alternate","or equal not allowed","no or equal"])
        yes_sub = any(p in t for p in [
            "substitution allowed","substitutions allowed","allow substitutions","allowed substitutions",
            "alternates allowed","alternate allowed","alternatives allowed","brand name or equal",
            "approved equal","approved equivalent","approved alternate","approved alternative",
            "or-equal","or equal","preapproved equal","pre-approved equal"
        ]) or bool(re.search(r"or\s+(approved\s+)?(equal|equivalent|alternate|alternative)s?", t))

        if install and yes_sub and (supply or True):
            detected.append("Lighting (Supply + Installation + Substitution Allowed)")
        elif supply and install and no_sub:
            detected.append("Lighting (Supply + Installation + Substitution Not allowed)")
        elif supply and yes_sub and not install:
            detected.append("Lighting (Supply + Substitution Allowed)")
        elif install:
            detected.append("Lighting (Only Installation)")

    if any(k in t for k in ["hvac","rtu","air handler","chiller","boiler"]):
        detected.append("HVAC")
    if any(k in t for k in ["solar","photovoltaic","pv system","pv array"]):
        detected.append("Solar PV")

    # dedupe
    return list(dict.fromkeys(detected)) or ["General Construction"]

# ----------------- POINTS -----------------
def compute_points_table_rows(rfp_text: str):
    profiles = detect_work_profiles(rfp_text)
    project_state = extract_project_state_simple(rfp_text)
    recommended = set()
    for p in profiles:
        recommended.update(PROFILE_TO_COMPANIES.get(p, []))

    rows = []
    for comp in COMPANY_LOCATIONS.keys():
        loc = get_company_base_location(comp) or COMPANY_LOCATIONS.get(comp, ", ")
        comp_state_raw = get_company_base_state(comp)
        comp_state_abbr = normalize_state_for_comparison(comp_state_raw)
        license_states = COMPANY_LICENSE_STATES.get(comp, set())
        base = 5 if comp in recommended else 0
        bonus = 0
        if base > 0 and project_state:
            # Normalize project_state (already 2-letter) and compare
            if project_state == comp_state_abbr or project_state in license_states:
                bonus = 5
        rows.append({
            "Company Name": comp,
            "Base Points": base,
            "State Bonus": bonus,
            "Total Points": base + bonus,
        })

    rows = sorted(rows, key=lambda r: (-int(r["Total Points"]), r["Company Name"]))
    allowed = [r["Company Name"] for r in rows if r["Total Points"] > 0] or [r["Company Name"] for r in rows]
    return profiles, project_state, rows, allowed

# ----------------- RULE-BASED SCORING (DB-driven) -----------------
QUESTIONS_A_I = [
    ("a", "Is project state and company state same?"),
    ("b", "Is the company or its subcontractor licensed in the project state (if required/mandatory)?"),
    ("c", "If site investigation/visit required/mandatory, can company/subcontractor do this?"),
    ("d", "Is company capable of fulfilling BABA/BAA/Davis Bacon (if required/mandatory)?"),
    ("e", "Is company capable of fulfilling specific qualifications for the project (if mandatory/required)?"),
    ("f", "Can companies meet SBE, MBE, WBE, HUB goals (if required)?"),
    ("g", "Can company meet specific security clearance or working hour restrictions (if required)?"),
    ("h", "Can company provide bond (payment/performance/bid) (if required)?"),
    ("i", "Can company provide insurance (if required)?"),
]

def _llm_call(client, model: str, system_prompt: str, user_prompt: str):
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt.strip()},
                {"role": "user", "content": user_prompt.strip()},
            ],
            temperature=0,
        )
        text = resp.choices[0].message.content
        usage = getattr(resp, "usage", None)
        request_id = getattr(resp, "id", None)
        return text, usage, request_id
    except BaseException as e:
        raise e

# ----------------- CONCISE DB EXPORT SUMMARY -----------------
def _has_any(text: str, keys: list[str]) -> bool:
    t = (text or "").lower()
    return any(k in t for k in keys)

def _build_concise_export(
    rfp_text: str,
    profiles: list[str],
    project_state: str | None,
    allowed_companies: list[str],
    company_scores: dict,
    bypass_decision: bool = False,
) -> dict:
    # Always choose a single best company even when totals tie.
    top_company = _select_best_company(allowed_companies, company_scores, points_rows=None)
    top_total = int((company_scores.get(top_company) or {}).get("total") or -1) if top_company else -1
    questions_count = len(QUESTIONS_A_I) or 9
    percent = max(0, min(100, round(100 * (top_total / (questions_count * 10))) if top_total >= 0 else 0))

    # Simple requirement flags from RFP
    t = (rfp_text or "").lower()
    site = _has_any(t, ["site visit", "site investigation", "pre-bid meeting", "walkthrough"]) \
        and "Site visit required" or None
    baba = _has_any(t, ["baba", "build america", "baa", "buy american", "davis bacon"]) \
        and "BABA/BAA/Davis Bacon" or None
    bonds = _has_any(t, ["bid bond", "performance bond", "payment bond"]) and "Bonds" or None
    ins = _has_any(t, ["insurance", "general liability", "workers compensation"]) and "Insurance" or None
    reqs = ", ".join([x for x in [site, baba, bonds, ins] if x]) or "No special requirements detected"

    short_type = (profiles[0] if profiles else "Project")
    state = project_state or "Unknown"
    decision = (company_scores.get(top_company, {}).get("decision") if top_company else None) or "No-Go"
    if bypass_decision:
        decision = "EVALUATED"

    if bypass_decision:
        lines = [
            f"RFP for {short_type} in {state}. Full evaluation generated (no Go/No-Go decision enforced).",
            f"Highest score observed: {top_company or 'N/A'} ({top_total}/90).",
            f"Key requirements noted: {reqs}.",
        ]
    else:
        lines = [
            f"RFP for {short_type} in {state}. Top company: {top_company or 'N/A'} ({top_total}/90), decision {decision}.",
            f"Key requirements noted: {reqs}.",
        ]
    summary = " ".join(lines)

    return {
        "b_name": "Not Found",
        "due_date": "Not Found",
        "state": state,
        "scope": short_type,
        "type": short_type,
        "scoring": percent,
        "comp_name": ("ALL" if bypass_decision else (top_company or "Not Found")),
        "decision": decision,
        "summary": summary,
    }

def _concat_company_db_text(sections: dict) -> str:
    parts: list[str] = []
    d = sections.get("details") or {}
    p = sections.get("preferences") or {}
    caps = sections.get("capabilities") or []
    web_ctx = sections.get("web_context") or ""
    case_studies = sections.get("case_studies") or []
    extra_tables = sections.get("extra_tables") or {}
    if d:
        parts.append("DETAILS: " + " ".join(str(d.get(k, "")) for k in ["address", "website"]))
    if p:
        parts.append("PREFERENCES: " + " ".join(str(p.get(k, "")) for k in ["deal_breakers", "deal_makers", "preferred_states"]))
    for c in caps:
        parts.append("CAPABILITY: " + " ".join([
            str(c.get("capability_title", "")),
            str(c.get("capability_description", "")),
            str(c.get("naics_codes", "")),
        ]))
    if web_ctx:
        parts.append("WEB_CONTEXT: " + str(web_ctx)[:8000])
    if case_studies:
        for cs in case_studies:
            if isinstance(cs, dict):
                items = []
                for k, v in cs.items():
                    kl = str(k).lower()
                    if kl in {"id", "company_id", "companyid", "company_name", "created_at", "updated_at"}:
                        continue
                    if v is None:
                        continue
                    s = str(v).strip()
                    if not s:
                        continue
                    items.append(f"{k}: {s}")
                if items:
                    parts.append("CASE_STUDY: " + " | ".join(items)[:8000])
            else:
                s = str(cs).strip()
                if s:
                    parts.append("CASE_STUDY: " + s[:8000])
    # Generic extra tables (e.g., past_performance, company_performance, case_study_* child tables)
    try:
        if isinstance(extra_tables, dict) and extra_tables:
            drop = {
                "id", "company_id", "companyid", "company_name", "created_at", "updated_at",
                "case_study_id", "casestudy_id",
            }
            # cap overall appended context to avoid runaway token growth
            total_chars = 0
            char_budget = 24000
            for tbl, rows in extra_tables.items():
                if not rows or not tbl:
                    continue
                label = str(tbl).upper()
                # cap per-table rows to keep context tight
                for r in (rows[:25] if isinstance(rows, list) else []):
                    if not isinstance(r, dict):
                        continue
                    kv = _row_to_compact_kv(r, drop_keys_lower=drop, max_field_chars=600)
                    if not kv:
                        continue
                    line = f"{label}: {kv}"
                    total_chars += len(line)
                    if total_chars > char_budget:
                        break
                    parts.append(line)
                if total_chars > char_budget:
                    break
    except Exception:
        pass

    return "\n".join([s for s in parts if s])

def _find_keyword_evidence(text: str, keywords: list[str], window: int = 120) -> dict | None:
    """Return first/earliest matched keyword + a nearby snippet for remarks."""
    t = text or ""
    low = t.lower()
    best_i = None
    best_kw = None
    for kw in (keywords or []):
        k = (kw or "").strip().lower()
        if not k:
            continue
        i = low.find(k)
        if i != -1 and (best_i is None or i < best_i):
            best_i = i
            best_kw = kw
    if best_i is None or best_kw is None:
        return None
    start = max(0, int(best_i) - window)
    end = min(len(t), int(best_i) + len(str(best_kw)) + window)
    snippet = t[start:end].replace("\n", " ")
    snippet = re.sub(r"\s+", " ", snippet).strip()
    return {"keyword": best_kw, "snippet": snippet}

# ----------------- WEBSITE SCRAPING -----------------
def _extract_visible_text(html: str) -> str:
    try:
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = soup.get_text(separator=" ")
        text = re.sub(r"\s+", " ", text).strip()
        return text
    except Exception:
        return ""

def _same_domain(url: str, base: str) -> bool:
    from urllib.parse import urlparse
    try:
        u1 = urlparse(url)
        u2 = urlparse(base)
        return (u1.netloc or "").lower() == (u2.netloc or "").lower()
    except Exception:
        return False

def _normalize_link(href: str, base: str) -> str | None:
    from urllib.parse import urljoin
    if not href:
        return None
    href = href.strip()
    if href.startswith("mailto:") or href.startswith("tel:"):
        return None
    try:
        return urljoin(base, href)
    except Exception:
        return None

def _scrape_site(start_url: str, max_pages: int = 5, max_chars: int = 120000, ignore_ssl: bool = False, use_sitemap: bool = True) -> str:
    if not _SCRAPE_OK:
        return ""
    seen: set[str] = set()
    queue: list[str] = [start_url]
    collected: list[str] = []
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Connection": "keep-alive",
    }
    # Optionally seed queue with sitemap URLs (helps sites with minimal navigation)
    if use_sitemap:
        try:
            from urllib.parse import urlparse, urljoin
            parsed = urlparse(start_url)
            root = f"{parsed.scheme}://{parsed.netloc}"
            sm_url = urljoin(root, "/sitemap.xml")
            try:
                sm_resp = requests.get(sm_url, timeout=15, headers=headers, verify=(not ignore_ssl))
            except Exception:
                sm_resp = requests.get(sm_url, timeout=15, headers=headers, verify=False)
            if sm_resp.status_code < 400 and len(queue) < max_pages:
                try:
                    sm_soup = BeautifulSoup(sm_resp.text, "xml")
                    locs = [loc.get_text().strip() for loc in sm_soup.find_all("loc")]
                    for l in locs[: max(0, max_pages * 2)]:
                        if l and l not in queue:
                            queue.append(l)
                except Exception:
                    pass
        except Exception:
            pass

    while queue and len(seen) < max_pages and sum(len(x) for x in collected) < max_chars:
        url = queue.pop(0)
        if url in seen:
            continue
        seen.add(url)
        try:
            try:
                resp = requests.get(url, timeout=20, headers=headers, verify=(not ignore_ssl))
            except Exception:
                if not ignore_ssl:
                    resp = requests.get(url, timeout=20, headers=headers, verify=False)
                else:
                    raise
            if resp.status_code >= 400:
                continue
            text = _extract_visible_text(resp.text)
            if text:
                collected.append(text)
            # discover a few links from this page
            try:
                soup = BeautifulSoup(resp.text, "html.parser")
                links = [a.get("href") for a in soup.find_all("a", href=True)]
                for href in links:
                    full = _normalize_link(href, url)
                    if not full:
                        continue
                    if not _same_domain(full, start_url):
                        continue
                    if full not in seen and len(queue) < max_pages * 3:
                        # prioritize common info pages
                        low = href.lower()
                        if any(k in low for k in ["about", "service", "capabil", "project", "portfolio", "qualification", "safety", "license", "compliance", "contact", "who-we-are", "what-we-do", "experience"]):
                            queue.insert(0, full)
                        else:
                            queue.append(full)
            except Exception:
                pass
        except Exception:
            continue
    return "\n\n".join(collected)[:max_chars]

def ensure_company_web_context(company_names: List[str], refresh: bool = False, max_pages: int = 5, max_chars: int = 120000, ignore_ssl: bool = False, use_sitemap: bool = True) -> Dict[str, str]:
    results: Dict[str, str] = {}
    ok, conn, _, _ = _open_mysql_or_create()
    if not ok:
        return results
    _ensure_company_tables(conn)
    _ensure_scrape_tables(conn)
    try:
        details_by_name: Dict[str, dict] = {}
        with conn.cursor() as c:
            c.execute("SELECT company_name, website FROM company_details WHERE website IS NOT NULL AND website != ''")
            for r in c.fetchall():
                details_by_name[r["company_name"]] = r
        existing: Dict[str, dict] = {}
        with conn.cursor() as c:
            c.execute("SELECT company_name, url, content FROM company_web_context")
            for r in c.fetchall():
                existing[r["company_name"]] = r
        # helper to update website in company_details if missing
        def _update_company_website(name: str, site: str):
            try:
                with conn.cursor() as c:
                    c.execute("UPDATE company_details SET website=%s WHERE company_name=%s", (site, name))
                conn.commit()
            except Exception:
                pass

        for name in company_names:
            # Always prefer mapped website; fall back to DB if no mapping
            mapped = _get_mapped_website(name)
            db_site = (details_by_name.get(name, {}) or {}).get("website")
            chosen = _normalize_url(mapped or db_site)
            if chosen and chosen != _normalize_url(db_site):
                _update_company_website(name, chosen)
            site = chosen
            if not site:
                continue
            have = existing.get(name)
            if have and not refresh and (have.get("content") or "").strip():
                results[name] = have.get("content") or ""
                continue
            text = _scrape_site(site, max_pages=max_pages, max_chars=max_chars, ignore_ssl=ignore_ssl, use_sitemap=use_sitemap)
            results[name] = text
            with conn.cursor() as c:
                c.execute(
                    (
                        "INSERT INTO company_web_context (company_name, url, content) "
                        "VALUES (%s,%s,%s) "
                        "ON DUPLICATE KEY UPDATE url=VALUES(url), content=VALUES(content)"
                    ),
                    (name, site, text),
                )
            conn.commit()
    finally:
        try:
            conn.close()
        except Exception:
            pass
    return results

def refresh_all_company_web_context(refresh: bool = True, max_pages: int = 50, max_chars: int = 300000, ignore_ssl: bool = False, use_sitemap: bool = True) -> Dict[str, str]:
    # Canonical company names used in DB
    target_companies = ["Ikio Led Lighting LLC", "METCO Engineering, Inc.", "Sunsprint Engineering", "METCO Engineering, Inc.", "SUNSPRINT ENGINEERING"]
    # Deduplicate with casing preserved for DB
    # Prefer DB variants first
    ordered = []
    seen = set()
    for n in ["Ikio Led Lighting LLC", "METCO Engineering, Inc.", "Sunsprint Engineering"]:
        if n not in seen:
            ordered.append(n)
            seen.add(n)
    return ensure_company_web_context(ordered, refresh=refresh, max_pages=max_pages, max_chars=max_chars, ignore_ssl=ignore_ssl, use_sitemap=use_sitemap)

def _missing_web_context(companies: List[str]) -> List[str]:
    ok, conn, _, _ = _open_mysql_or_create()
    if not ok:
        return companies
    _ensure_scrape_tables(conn)
    existing = set()
    try:
        with conn.cursor() as c:
            c.execute("SELECT company_name FROM company_web_context WHERE content IS NOT NULL AND content != ''")
            for r in c.fetchall():
                existing.add(r.get("company_name"))
    finally:
        try:
            conn.close()
        except Exception:
            pass
    out = []
    for n in companies:
        if n not in existing:
            out.append(n)
    return out

def _contains_any(text: str, keywords: list[str]) -> bool:
    t = (text or "").lower()
    return any(k in t for k in keywords)

def _score_from_bool(val: bool | None) -> int:
    if val is True:
        return 10
    if val is False:
        return 3
    return 6

def rule_score_company(
    company_name: str,
    sections: dict,
    detected_state: str | None,
    profiles: list[str],
    bypass_decision: bool = False,
):
    text_blob = _concat_company_db_text(sections)
    base_state = get_company_base_state(company_name)
    base_state_abbr = normalize_state_for_comparison(base_state)
    license_states = COMPANY_LICENSE_STATES.get(company_name, set())

    # a) state match (normalize both for comparison)
    a_ok = (detected_state is not None and base_state_abbr == detected_state)

    # b) license in project state
    b_ok = (detected_state is not None and detected_state in license_states)

    # c) site investigation capability
    c_keys = [
        "site investigation", "site visit", "field survey", "walkthrough", "pre-bid meeting",
        "job walk", "prebid", "site survey", "field verification", "onsite assessment", "site assessment"
    ]
    c_ev = _find_keyword_evidence(text_blob, c_keys)
    if not c_ev:
        hits = rag_retrieve(text_blob, "site visit / pre-bid walkthrough / field verification requirements", top_k=2)
        if hits:
            c_ev = {"keyword": "RAG", "snippet": f"(RAG {hits[0]['score']}) {hits[0]['text'][:240]}..."}
    c_ok = True if c_ev else None

    # d) procurement compliance
    d_keys = [
        "baba", "build america", "baa", "buy american", "buy america", "davis bacon", "prevailing wage",
        "american iron and steel", "ais", "domestic preference", "made in usa"
    ]
    d_ev = _find_keyword_evidence(text_blob, d_keys)
    if not d_ev:
        hits = rag_retrieve(text_blob, "procurement compliance requirements (BABA/Buy America/Davis-Bacon/AIS/domestic preference)", top_k=2)
        if hits:
            d_ev = {"keyword": "RAG", "snippet": f"(RAG {hits[0]['score']}) {hits[0]['text'][:240]}..."}
    d_ok = True if d_ev else None

    # e) qualifications (use profile keywords)
    qual_keys = [
        "experience", "references", "past performance", "qualified", "licensed", "certified",
        "238210", "electrical contractor", "lighting retrofit", "led retrofit", "lighting design",
        "engineer", "engineering", "master electrician"
    ]
    if any("Lighting" in p for p in profiles):
        qual_keys += ["lighting", "electrical", "retrofit", "led"]
    e_ev = _find_keyword_evidence(text_blob, qual_keys)
    if not e_ev:
        hits = rag_retrieve(text_blob, "qualifications / similar project experience / references / certifications / past performance", top_k=2)
        if hits:
            e_ev = {"keyword": "RAG", "snippet": f"(RAG {hits[0]['score']}) {hits[0]['text'][:240]}..."}
    e_ok = True if e_ev else None

    # f) SBE/MBE/WBE/HUB
    f_keys = [
        "sbe", "mbe", "wbe", "hub", "h.u.b.", "dbe", "db e", "8(a)", "hubzone",
        "minority-owned", "woman-owned", "women owned", "supplier diversity"
    ]
    f_ev = _find_keyword_evidence(text_blob, f_keys)
    if not f_ev:
        hits = rag_retrieve(text_blob, "SBE/MBE/WBE/DBE/HUB/8(a)/HUBZone certifications and supplier diversity", top_k=2)
        if hits:
            f_ev = {"keyword": "RAG", "snippet": f"(RAG {hits[0]['score']}) {hits[0]['text'][:240]}..."}
    f_ok = True if f_ev else None

    # g) security/working hours
    g_keys = [
        "background", "background check", "badge", "badging", "security", "cjis", "twic", "drug test",
        "after hours", "night work", "off-hours", "weekend work"
    ]
    g_ev = _find_keyword_evidence(text_blob, g_keys)
    if not g_ev:
        hits = rag_retrieve(text_blob, "security requirements (badging/background checks/CJIS/TWIC/drug tests) and off-hours work", top_k=2)
        if hits:
            g_ev = {"keyword": "RAG", "snippet": f"(RAG {hits[0]['score']}) {hits[0]['text'][:240]}..."}
    g_ok = True if g_ev else None

    # h) bonds
    h_keys = [
        "bond", "bonding", "surety", "bonded", "bonding capacity", "bid security", "performance bond", "payment bond"
    ]
    h_ev = _find_keyword_evidence(text_blob, h_keys)
    if not h_ev:
        hits = rag_retrieve(text_blob, "bonding requirements (bid/performance/payment bonds, surety, bonding capacity)", top_k=2)
        if hits:
            h_ev = {"keyword": "RAG", "snippet": f"(RAG {hits[0]['score']}) {hits[0]['text'][:240]}..."}
    h_ok = True if h_ev else None

    # i) insurance
    i_keys = [
        "insurance", "general liability", "cgl", "gl", "workers compensation", "workers' compensation", "wc",
        "auto liability", "umbrella", "excess liability", "coi"
    ]
    i_ev = _find_keyword_evidence(text_blob, i_keys)
    if not i_ev:
        hits = rag_retrieve(text_blob, "insurance requirements (general liability, workers comp, auto liability, umbrella/excess, COI)", top_k=2)
        if hits:
            i_ev = {"keyword": "RAG", "snippet": f"(RAG {hits[0]['score']}) {hits[0]['text'][:240]}..."}
    i_ok = True if i_ev else None

    q_text_map = {ltr: txt for (ltr, txt) in QUESTIONS_A_I}

    def mk_row(letter: str, ok: bool | None, ev: dict | None, no_fix: str, keys: list[str]):
        score = _score_from_bool(ok)
        if ok is True:
            rec = "-"
            if ev:
                remark = f"Matched '{ev.get('keyword')}' - {ev.get('snippet')}"
            else:
                remark = "Evidence found in company context."
        else:
            rec = no_fix
            sample = ", ".join(keys[:6]) + ("…" if len(keys) > 6 else "")
            remark = f"No evidence found for keywords ({sample}) in company context (details/preferences/capabilities/web/case studies)."
        return {"q": letter, "question": q_text_map.get(letter, ""), "score": score, "remark": remark, "recommendation": rec}

    rows = []
    # a) state match remark is computed (not constant)
    rows.append({
        "q": "a",
        "question": q_text_map.get("a", ""),
        "score": _score_from_bool(a_ok),
        "remark": (
            f"Project state={detected_state}; company base_state={base_state_abbr}."
            + (" Match." if a_ok else " No match.")
        ),
        "recommendation": "-" if a_ok else "Use local subcontractor or confirm ability to operate in-state",
    })
    # b) license remark is computed from detected_state + license list
    lic_list = ", ".join(sorted(list(license_states))) if license_states else "(none listed)"
    rows.append({
        "q": "b",
        "question": q_text_map.get("b", ""),
        "score": _score_from_bool(b_ok),
        "remark": f"Project state={detected_state}; licensed_states={lic_list}.",
        "recommendation": "-" if b_ok else "Engage state-licensed subcontractor",
    })
    # c-i) evidence-driven remarks (keyword + snippet)
    rows.extend([
        mk_row("c", c_ok, c_ev, "Schedule pre-bid/site visit; assign local team", c_keys),
        mk_row("d", d_ok, d_ev, "Source compliant materials; attach compliance plan", d_keys),
        mk_row("e", e_ok, e_ev, "Attach similar project references meeting quals", qual_keys),
        mk_row("f", f_ok, f_ev, "Partner with certified local firm", f_keys),
        mk_row("g", g_ok, g_ev, "Confirm background checks; plan off-hours work", g_keys),
        mk_row("h", h_ok, h_ev, "Confirm surety line; secure required bonds", h_keys),
        mk_row("i", i_ok, i_ev, "Confirm COI limits with carrier", i_keys),
    ])
    total = sum(r["score"] for r in rows)
    decision = ("Evaluated" if bypass_decision else ("Go" if total >= 70 else "No-Go"))
    return {"rows": rows, "total": total, "decision": decision, "usages": []}

# ----------------- EVALUATION SEQUENCE (Supply/Install/Substitution) -----------------
def analyze_work_criteria(rfp_text: str) -> dict:
    t = (rfp_text or "").lower()
    supply = any(k in t for k in ["supply", "furnish", "provide materials", "materials", "equipment", "procure", "purchase", "f&i", "furnish and install"])
    install = any(k in t for k in ["install", "installation", "replace", "retrofit", "mount", "erect"])
    no_sub = any(p in t for p in ["no substitution", "substitution not allowed", "no alternates", "no alternate", "or equal not allowed", "no or equal"])
    yes_sub = any(p in t for p in [
        "substitution allowed", "substitutions allowed", "allow substitutions", "allowed substitutions",
        "alternates allowed", "alternate allowed", "alternatives allowed", "brand name or equal",
        "approved equal", "approved equivalent", "approved alternate", "approved alternative",
        "or-equal", "or equal", "preapproved equal", "pre-approved equal"
    ]) or bool(re.search(r"or\s+(approved\s+)?(equal|equivalent|alternate|alternative)s?", t))
    return {
        "supply": supply,
        "installation": install,
        "substitution": ("Allowed" if yes_sub and not no_sub else ("Not allowed" if no_sub and not yes_sub else "Not specified")),
    }

def build_evaluation_sequence_md(criteria: dict) -> str:
    lines = [
        "### EVALUATION SEQUENCE (by criteria)",
        f"- Supply: {'Yes' if criteria.get('supply') else 'No'}",
        f"- Installation: {'Yes' if criteria.get('installation') else 'No'}",
        f"- Substitution: {criteria.get('substitution')}",
    ]
    return "\n".join(lines)

def order_companies_by_profile(profiles: list[str], allowed_companies: list[str]) -> list[str]:
    priority = [
        "Lighting (Supply + Installation + Substitution Allowed)",
        "Lighting (Supply + Installation + Substitution Not allowed)",
        "Lighting (Supply + Substitution Allowed)",
        "Lighting (Only Installation)",
        "HVAC",
        "Solar PV",
    ]
    for p in priority:
        if p in profiles:
            pri = [c for c in PROFILE_TO_COMPANIES.get(p, []) if c in allowed_companies]
            rest = [c for c in allowed_companies if c not in pri]
            return pri + rest
    return allowed_companies


def _select_best_company(
    allowed_companies: list[str],
    company_scores: dict,
    points_rows: list[dict] | None = None,
) -> str | None:
    """
    Select exactly ONE best company.
    Tie-breakers (in order):
    - total score (desc)
    - question 'e' (qualifications/experience) score (desc)
    - question 'a' (state match) score (desc)
    - question 'b' (license in state) score (desc)
    - count of 10-point criteria (desc)
    - points table total points (desc) if available
    - company name (asc) for stable determinism
    """
    if not allowed_companies or not company_scores:
        return None

    points_map: dict[str, int] = {}
    if points_rows:
        try:
            for r in points_rows:
                nm = r.get("Company Name")
                if nm:
                    points_map[str(nm)] = int(r.get("Total Points") or 0)
        except Exception:
            points_map = {}

    def _q_score(comp: str, q: str) -> int:
        try:
            rows = (company_scores.get(comp) or {}).get("rows") or []
            for rr in rows:
                if rr.get("q") == q:
                    return int(rr.get("score") or 0)
        except Exception:
            return 0
        return 0

    def _count_tens(comp: str) -> int:
        try:
            rows = (company_scores.get(comp) or {}).get("rows") or []
            return sum(1 for rr in rows if int(rr.get("score") or 0) >= 10)
        except Exception:
            return 0

    def sort_key(comp: str):
        total = int((company_scores.get(comp) or {}).get("total") or 0)
        e = _q_score(comp, "e")
        a = _q_score(comp, "a")
        b = _q_score(comp, "b")
        tens = _count_tens(comp)
        pts = int(points_map.get(comp, 0))
        return (-total, -e, -a, -b, -tens, -pts, str(comp))

    candidates = [c for c in allowed_companies if c in company_scores]
    if not candidates:
        return None
    return sorted(candidates, key=sort_key)[0]

# ----------------- FINAL ANALYSIS (uses per-question API) -----------------
def analyze_with_gpt(
    rfp_text: str,
    file_name: str | None = None,
    bypass_decision: bool = False,
    file_hash: str | None = None,
    attached_files: list[tuple[str, bytes]] | None = None,
):
    # 1) Deterministic context
    profiles, project_state, points_rows, allowed_companies = compute_points_table_rows(rfp_text)
    rfp_excerpt = rfp_text[:12000]
    db_ctx = load_company_sections_from_db(allowed_companies)
    db_block = format_company_db_context(db_ctx)
    criteria = analyze_work_criteria(rfp_text)
    evaluation_sequence_md = build_evaluation_sequence_md(criteria)
    ordered_companies = order_companies_by_profile(profiles, allowed_companies)

    # 1a) Ensure website context for each company and inject into context
    try:
        scraped = ensure_company_web_context(allowed_companies, refresh=False, max_pages=5)
        if scraped:
            for comp, text in scraped.items():
                if comp in db_ctx:
                    db_ctx[comp]["web_context"] = text
    except Exception:
        pass

    # points md
    points_md = "\n".join([f"| {r['Company Name']} | {r['Base Points']} | {r['State Bonus']} | {r['Total Points']} |"
                           for r in points_rows])
    points_table = f"| Company | Base | Bonus | Total |\n|---|---:|---:|---:|\n{points_md}"

    # 2) RULE-BASED scoring from DB (no per-question model calls)
    company_scores = {}
    for comp in allowed_companies:
        sections = db_ctx.get(comp) or {}
        company_scores[comp] = rule_score_company(comp, sections, project_state, profiles, bypass_decision=bypass_decision)
    usage_log: list[dict] = []

    # Always select exactly one best company (even if totals tie).
    best_company = _select_best_company(allowed_companies, company_scores, points_rows=points_rows)

    # Enforce single-company decision when not bypassing.
    # Best company is "Go" only if it meets the threshold; all others are "No-Go".
    if (not bypass_decision) and company_scores and best_company:
        best_total = int((company_scores.get(best_company) or {}).get("total") or 0)
        best_decision = "Go" if best_total >= 70 else "No-Go"
        for c in allowed_companies:
            if c in company_scores:
                company_scores[c]["decision"] = best_decision if c == best_company else "No-Go"

    # Build deterministic markdown tables for all companies
    detailed_tables_lines: list[str] = []
    for comp in ordered_companies:
        cs = company_scores[comp]
        detailed_tables_lines.append(f"### {comp}")
        detailed_tables_lines.append("| Question | Score | Remark | Recommendation |")
        detailed_tables_lines.append("|---|---:|---|---|")
        for row in cs["rows"]:
            qtxt = row.get('question') or row.get('q')
            detailed_tables_lines.append(f"| {qtxt} | {row['score']} | {row['remark']} | {row['recommendation']} |")
        detailed_tables_lines.append(f"Total Score: {cs['total']} - Decision: {cs['decision']}")
        detailed_tables_lines.append("")
    detailed_tables_md = "\n".join(detailed_tables_lines)

    # 4) Final report assembly call - sanitize non-serializable fields
    final_sys = "You are a Senior EPC Bid Evaluation Specialist. Format the final report clearly."
    import copy
    clean_payload = copy.deepcopy(company_scores)
    for comp in clean_payload:
        if "usages" in clean_payload[comp]:
            for u in clean_payload[comp]["usages"]:
                if isinstance(u.get("usage"), object):
                    u["usage"] = str(u.get("usage"))
                if isinstance(u.get("request_id"), object):
                    u["request_id"] = str(u.get("request_id"))

    final_payload = {
        "detected_profiles": profiles,
        "detected_state": project_state,
        "allowed_companies": ordered_companies,
        "best_company": best_company,
        "points_table_md": points_table,
        "company_scores": clean_payload,
        "detailed_tables_md": detailed_tables_md,
        "evaluation_sequence_md": evaluation_sequence_md,
    }
    final_usr = f"""
RFP EXCERPT (for quoting compliance answers verbatim if needed):
{rfp_excerpt[:9000]}

STRUCTURED DATA:
{json.dumps(final_payload, indent=2, ensure_ascii=False)}

BYPASS_DECISION: {str(bool(bypass_decision))}

RENDER the final response with these sections:
1) HEADER
2) STEP 1 - RFP DATA EXTRACTION (table)
   Include a row named "Submission Method" and normalize its value using ONLY these options (list multiple if allowed): 
   1) Electronically, 2) Email, 3) Submission via link, 4) Bids return to Email ID, 5) Sealed Bid, 6) Hand Delivery, 7) Mailed Bid. 
   Search for cues such as: "electronic submission", "submit electronically", "via portal", "online portal", "upload", "email to", "send to email", 
   "submission link", "apply at link", "return to <email>", "sealed bid", "sealed envelope", "hand deliver", "deliver in person", "mail", "mailed bids", "USPS", "FedEx", "UPS". 
   If nothing is found, write "Not specified". Provide a brief quote/citation from the RFP where available.
3) COMPANY SELECTION CRITERIA (show points_table_md as Markdown)
3a) EVALUATION SEQUENCE (Supply/Installation/Substitution)
4) DETAILED BID EVALUATION (for each company, show a..i with the exact score and short remark provided)
5) COMPLIANCE TABLE (quote from the RFP excerpt)
6) FINAL RECOMMENDATION - Write ~180–250 words recommending ONLY the single best company (use best_company from STRUCTURED DATA). Justify using a..i (licensing/state match, site visit, BABA/BAA/DB, qualifications, SBE/MBE/WBE/HUB, security/working hours, bonds, insurance); include key risks with mitigations; next steps; end with one‑line executive takeaway.
   IMPORTANT: If bypass_decision is enabled, DO NOT label any company as GO/NO-GO; just provide a ranked assessment.
6a) SCOPE OF WORK - Immediately after Final Recommendation, add a concise paragraph or short bullets summarizing the technical scope (work type, delivery mode: supply/installation/substitution, and key requirements like site visit, BABA/BAA/DB, bonds, insurance) based on the RFP excerpt.
7) DATABASE EXPORT JSON (keys: b_name, due_date, state, scope, type, scoring, comp_name, decision, summary)
"""
    # Use LLM only for final formatting/narrative (optional)
    cli = _get_openai_client()
    final_text = None
    if cli:
        try:
            final_text, _, _ = _llm_call(cli, _get_selected_model(), final_sys, final_usr)
        except Exception:
            final_text = None
    if not final_text:
        # Minimal fallback formatting without LLM
        sow_fb = build_scope_of_work_fallback(rfp_text, profiles)
        out_lines = [
            "Project Name: Not Found",
            "",
            "### STEP 1 - RFP DATA EXTRACTION",
            "(Not extracted in fallback mode)",
            "",
            "### COMPANY SELECTION CRITERIA",
            points_table,
            "",
            evaluation_sequence_md,
            "",
            "### SCOPE OF WORK",
            sow_fb,
        ]
        final_text = "\n".join(out_lines)
    else:
        # Append evaluation sequence only (tables removed as requested)
        final_text = f"{final_text}\n\n{evaluation_sequence_md}"

    # 5) Show usage summary so you SEE API consumption
    if usage_log:
        total_prompt = sum((u.get("prompt_tokens") or 0) for u in usage_log)
        total_completion = sum((u.get("completion_tokens") or 0) for u in usage_log)
        total_tokens = sum((u.get("total_tokens") or 0) for u in usage_log)
        _st_safe(
            st.info,
            f"🔎 OpenAI per-question scoring calls: {len(usage_log)} - "
            f"Prompt tokens: {total_prompt}, Completion tokens: {total_completion}, Total: {total_tokens}"
        )
        with st.expander("Show detailed API usage (per question)", expanded=False):
            for u in usage_log:
                _st_safe(
                    st.caption,
                    f"• {u['company']} Q{u['q']} | model={u['model']} | request_id={u['request_id']} | "
                    f"pt={u['prompt_tokens']} ct={u['completion_tokens']} tt={u['total_tokens']}"
                )

    # 6) Prefer any valid JSON object in the output (e.g., user/LLM payload),
    # then a labeled DB_EXPORT_JSON block; else build fallback
    top_json = parse_first_json_object(final_text)
    llm_export = parse_db_export_line(final_text)
    export = top_json or llm_export or _build_concise_export(
        rfp_text, profiles, project_state, allowed_companies, company_scores, bypass_decision=bypass_decision
    )
    # Populate export summary with required narrative sections for DB/UI:
    # - Context (Compliance Table)
    # - Final Recommendation
    # - Scope of Work
    try:
        parts: list[str] = []

        step1_text = extract_step1_data_extraction_section(final_text)
        if step1_text:
            parts.append("STEP 1 - RFP DATA EXTRACTION\n" + step1_text.strip())

        seq_text = extract_evaluation_sequence_section(final_text)
        if seq_text:
            parts.append("EVALUATION SEQUENCE\n" + seq_text.strip())

        eval_text = extract_detailed_bid_evaluation_section(final_text)
        if eval_text:
            parts.append("DETAILED BID EVALUATION\n" + eval_text.strip())

        ctx_text = extract_compliance_table_section(final_text)
        if ctx_text:
            parts.append("COMPLIANCE TABLE\n" + ctx_text.strip())

        fr_text = extract_final_recommendation_summary(final_text)
        if fr_text:
            parts.append("FINAL RECOMMENDATION\n" + fr_text.strip())

        sow_text = extract_scope_of_work_section(final_text)
        if not sow_text:
            sow_text = build_scope_of_work_fallback(rfp_text, profiles)
        if sow_text:
            parts.append("SCOPE OF WORK\n" + sow_text.strip())

        if parts:
            export["summary"] = "\n\n".join(parts).strip()
    except Exception:
        pass

    # If bypass is enabled, ensure DB export reflects "full evaluation" rather than a single-company decision.
    if bypass_decision and export:
        export["comp_name"] = "ALL"
        export["decision"] = "EVALUATED"
    if not llm_export:
        final_text = f"{final_text}\n\nDB_EXPORT_JSON: {json.dumps(export, ensure_ascii=False)}"
    ok1, msg1 = save_full_result_to_db(file_name, final_text, export, file_hash=file_hash)
    if ok1:
        _st_safe(st.success, "✅ Stored full evaluation output in database (evaluation_runs).")
    else:
        _st_safe(st.warning, f"⚠️ Could not store full output: {msg1}")
    bid_id = None
    if export:
        ok2, msg2, bid_id = save_bid_result_to_db(export)
        if ok2:
            if "duplicate" in (msg2 or "").lower():
                _st_safe(st.info, f"INFO: BID already existed; updated DB record ({msg2}).")
            else:
                _st_safe(st.success, "Saved structured BID export to database (bid_incoming).")
        else:
            _st_safe(st.warning, f"Could not save BID export: {msg2}")

    try:
        if bid_id and attached_files:
            ok_db, conn, _, _ = _open_mysql_or_create()
            if ok_db and conn:
                try:
                    _attach_uploaded_files_to_bid(conn, int(bid_id), attached_files)
                finally:
                    try:
                        conn.close()
                    except Exception:
                        pass
    except Exception:
        pass

    return final_text, bid_id

# ----------------- DUPLICATE DETECTION & RESULT RETRIEVAL -----------------
def get_file_hash(file_bytes: bytes) -> str:
    """Calculate SHA256 hash of file bytes."""
    import hashlib
    return hashlib.sha256(file_bytes).hexdigest()

def check_file_analyzed(file_hash: str) -> tuple[bool, dict | None]:
    """
    Check if a file (by hash) has already been analyzed.
    Returns (is_analyzed, file_info_dict or None).
    """
    ok, conn, _, _ = _open_mysql_or_create()
    if not ok:
        return False, None
    try:
        _ensure_eval_tables(conn)
        with conn.cursor() as c:
            # Check if file exists in uploaded_rfp_files
            c.execute(
                """
                SELECT original_filename
                FROM uploaded_rfp_files
                WHERE file_hash = %s
                ORDER BY uploaded_at DESC
                LIMIT 1
                """,
                (file_hash,)
            )
            file_row = c.fetchone()
            if not file_row:
                return False, None
            
            original_filename = file_row.get("original_filename")
            
            # Check if evaluation exists for this file
            c.execute(
                """
                SELECT id, file_name, result_text, export_json, created_at
                FROM evaluation_runs
                WHERE file_hash = %s
                ORDER BY created_at DESC
                LIMIT 1
                """,
                (file_hash,)
            )
            eval_row = c.fetchone()
            
            if eval_row and eval_row.get("result_text"):
                # File exists AND has evaluation results - return immediately
                return True, {
                    "file_name": original_filename,
                    "result_text": eval_row.get("result_text"),
                    "export_json": eval_row.get("export_json"),
                    "created_at": eval_row.get("created_at"),
                    "eval_id": eval_row.get("id"),
                    "file_hash": file_hash,  # Include hash for reference
                }
            # File exists in DB but no evaluation results yet
            return False, {"file_name": original_filename, "file_hash": file_hash}
    except Exception as e:
        return False, None
    finally:
        try:
            conn.close()
        except Exception:
            pass

def get_previous_evaluation_result(file_hash: str) -> str | None:
    """
    Retrieve previous evaluation result text for a file (by hash).
    Returns result_text if found, None otherwise.
    """
    is_analyzed, info = check_file_analyzed(file_hash)
    if is_analyzed and info and info.get("result_text"):
        return info.get("result_text")
    return None

# Added batch evaluation duplicate check
def check_batch_analyzed(batch_label: str) -> tuple[bool, dict | None]:
    """
    Check if a batch/folder evaluation has already been analyzed by looking up the
    evaluation_runs table by file_name (which stores the batch/folder label).

    Returns a tuple (is_analyzed, info_dict or None). If a previous evaluation exists,
    info_dict will contain result_text and other metadata.
    """
    if not batch_label:
        return False, None
    ok, conn, _, _ = _open_mysql_or_create()
    if not ok:
        return False, None
    try:
        # Ensure table exists
        _ensure_eval_tables(conn)
        with conn.cursor() as c:
            # Look up the most recent evaluation_run with this file_name
            c.execute(
                """
                SELECT id, file_name, result_text, export_json, created_at
                FROM evaluation_runs
                WHERE file_name = %s
                ORDER BY created_at DESC
                LIMIT 1
                """,
                (batch_label,),
            )
            row = c.fetchone()
            if row and row.get("result_text"):
                return True, {
                    "file_name": row.get("file_name"),
                    "result_text": row.get("result_text"),
                    "export_json": row.get("export_json"),
                    "created_at": row.get("created_at"),
                    "eval_id": row.get("id"),
                }
            return False, None
    except Exception:
        return False, None

# ----------------- ONEDRIVE AUTO SYNC -----------------
_ONEDRIVE_ALLOWED_EXTS = (".pdf", ".txt", ".docx", ".jpg", ".jpeg", ".png", ".tiff", ".tif", ".xlsx", ".xls")
_ONEDRIVE_SCHEDULER_STARTED = False
_ONEDRIVE_SCHEDULER_LOCK = threading.Lock()
_ONEDRIVE_SCAN_LOCK = threading.Lock()

def _parse_scan_times(times_str: str | None) -> list[tuple[int, int]]:
    raw = (times_str or "").strip()
    if not raw:
        return [(0, 0), (12, 0)]
    out: list[tuple[int, int]] = []
    for part in raw.split(","):
        p = part.strip()
        if not p:
            continue
        if ":" not in p:
            continue
        hh, mm = p.split(":", 1)
        try:
            h = int(hh)
            m = int(mm)
        except Exception:
            continue
        if 0 <= h <= 23 and 0 <= m <= 59:
            out.append((h, m))
    if not out:
        return [(0, 0), (12, 0)]
    return out

def _seconds_until_next_run(times: list[tuple[int, int]]) -> float:
    now = datetime.now()
    next_times: list[datetime] = []
    for h, m in (times or []):
        t = now.replace(hour=h, minute=m, second=0, microsecond=0)
        if t <= now:
            t = t + timedelta(days=1)
        next_times.append(t)
    if not next_times:
        # default to 12 hours
        return 12 * 3600
    nxt = min(next_times)
    return max(1.0, (nxt - now).total_seconds())

def _get_onedrive_root() -> str:
    try:
        override = st.session_state.get("onedrive_root_override")
    except Exception:
        override = None
    env_root = (
        os.getenv("ONEDRIVE_SCAN_ROOT", "").strip()
        or os.getenv("ONEDRIVE_SYNC_PATH", "").strip()
        or os.getenv("ONEDRIVE_ROOT", "").strip()
    )
    return (override or env_root or "").strip()

def _get_onedrive_scan_enabled() -> bool:
    try:
        override = st.session_state.get("onedrive_scan_enabled_override")
    except Exception:
        override = None
    if override is not None:
        return bool(override)
    env_val = os.getenv("ONEDRIVE_SCAN_ENABLED", "0").strip().lower()
    return env_val in ("1", "true", "yes", "y", "on")

def _get_onedrive_scan_times() -> list[tuple[int, int]]:
    try:
        override = st.session_state.get("onedrive_scan_times_override")
    except Exception:
        override = None
    return _parse_scan_times(override or os.getenv("ONEDRIVE_SCAN_TIMES", "00:00,12:00"))

def _get_onedrive_max_chars() -> int:
    try:
        override = st.session_state.get("onedrive_max_chars_override")
        if override is not None:
            return int(override)
    except Exception:
        pass
    try:
        return int(os.getenv("ONEDRIVE_MAX_CHARS", "250000"))
    except Exception:
        return 250000

def _list_project_folders(root_path: str) -> list[str]:
    try:
        entries = [os.path.join(root_path, n) for n in os.listdir(root_path)]
    except Exception:
        return []
    folders = [p for p in entries if os.path.isdir(p)]
    # If no subfolders, treat root as a single project folder
    if not folders:
        return [root_path]
    return sorted(folders)

def _iter_folder_files(folder_path: str) -> list[str]:
    files: list[str] = []
    for dirpath, _, filenames in os.walk(folder_path):
        for fname in filenames:
            if not fname:
                continue
            lf = fname.lower()
            if lf.endswith(_ONEDRIVE_ALLOWED_EXTS):
                files.append(os.path.join(dirpath, fname))
    return sorted(files)

def _folder_signature(folder_path: str) -> str:
    """Fast signature based on path, size, and mtime."""
    parts: list[str] = []
    for fp in _iter_folder_files(folder_path):
        try:
            stinfo = os.stat(fp)
            rel = os.path.relpath(fp, folder_path)
            parts.append(f"{rel}|{stinfo.st_size}|{int(stinfo.st_mtime)}")
        except Exception:
            continue
    joined = "||".join(sorted(parts))
    return hashlib.sha256(joined.encode("utf-8")).hexdigest()

def _folder_label_for_path(folder_path: str) -> str:
    base = os.path.basename(folder_path.rstrip("\\/")) or "PACKAGE"
    h = hashlib.sha1(folder_path.lower().encode("utf-8")).hexdigest()[:8]
    return f"FOLDER_{base}_{h}"

def _process_onedrive_folder(folder_path: str) -> bool:
    """Return True if evaluation ran, False if skipped."""
    log = _get_onedrive_logger()
    if not os.path.isdir(folder_path):
        log.warning("Skip: folder not found: %s", folder_path)
        return False
    # Prevent overlapping scans
    if not _ONEDRIVE_SCAN_LOCK.acquire(blocking=False):
        log.info("Skip: scan already running. folder=%s", folder_path)
        return False
    try:
        ok, conn, _, _ = _open_mysql_or_create()
        if not ok:
            log.error("DB connect failed; folder=%s", folder_path)
            return False
        try:
            _ensure_eval_tables(conn)
            _ensure_onedrive_tables(conn)
            signature = _folder_signature(folder_path)
            state = _get_onedrive_folder_state(conn, folder_path)
            last_sig = (state or {}).get("last_signature")
            label = _folder_label_for_path(folder_path)

            # If no changes, just mark scanned and skip.
            if last_sig and signature == last_sig:
                log.info("No changes; skipping. label=%s folder=%s", label, folder_path)
                _upsert_onedrive_folder_state(
                    conn,
                    folder_path,
                    folder_label=label,
                    signature=signature,
                    scanned=True,
                    evaluated=False,
                )
                return False

            files = _iter_folder_files(folder_path)
            if not files:
                log.info("No eligible files; skipping. label=%s folder=%s", label, folder_path)
                _upsert_onedrive_folder_state(
                    conn,
                    folder_path,
                    folder_label=label,
                    signature=signature,
                    scanned=True,
                    evaluated=False,
                )
                return False

            max_chars = _get_onedrive_max_chars()
            prefer_fast_pdf = True
            force_pdf_ocr = False
            pdf_fast_max_pages = 30

            combined_parts: list[str] = []
            attached_files: list[tuple[str, bytes]] = []
            combined_len = 0

            log.info("Processing folder: %s files=%d label=%s", folder_path, len(files), label)
            for fp in files:
                if combined_len >= int(max_chars):
                    log.info("Reached max_chars; stopping early. label=%s chars=%d", label, combined_len)
                    break
                try:
                    with open(fp, "rb") as fh:
                        fbytes = fh.read()
                    fname = os.path.basename(fp)
                    # Save file record (skip duplicates)
                    save_uploaded_file(fbytes, fname, skip_if_exists=True)
                    attached_files.append((fname, fbytes))
                    t = extract_text_any(
                        fbytes,
                        fname,
                        prefer_fast_pdf=prefer_fast_pdf,
                        force_pdf_ocr=force_pdf_ocr,
                        pdf_fast_max_pages=pdf_fast_max_pages,
                    ).strip()
                    if not t:
                        continue
                    block = f"\n\n===== FILE: {fname} =====\n{t}"
                    combined_parts.append(block)
                    combined_len += len(block)
                except Exception:
                    log.exception("Failed reading/extracting file: %s", fp)
                    continue

            combined_text = ("\n".join(combined_parts)).strip()
            if not combined_text:
                log.warning("No text extracted; skipping. label=%s folder=%s", label, folder_path)
                _upsert_onedrive_folder_state(
                    conn,
                    folder_path,
                    folder_label=label,
                    signature=signature,
                    scanned=True,
                    evaluated=False,
                )
                return False

            log.info("Evaluating folder: label=%s chars=%d", label, len(combined_text))
            _st_safe(st.info, f"OneDrive auto-sync: evaluating {label}")
            try:
                bypass = bool(st.session_state.get("bypass_decision"))
            except Exception:
                bypass = False
            _, _ = analyze_with_gpt(
                combined_text,
                file_name=label,
                bypass_decision=bypass,
                file_hash=get_file_hash(combined_text.encode("utf-8")),
                attached_files=attached_files,
            )
            log.info("Evaluation complete: label=%s", label)
            _upsert_onedrive_folder_state(
                conn,
                folder_path,
                folder_label=label,
                signature=signature,
                scanned=True,
                evaluated=True,
            )
            return True
        finally:
            try:
                conn.close()
            except Exception:
                pass
    finally:
        try:
            _ONEDRIVE_SCAN_LOCK.release()
        except Exception:
            pass

def scan_onedrive_root() -> int:
    """Scan the configured OneDrive root. Returns number of evaluated folders."""
    log = _get_onedrive_logger()
    root = _get_onedrive_root()
    if not root or not os.path.isdir(root):
        log.error("OneDrive root not found: %s", root)
        return 0
    folders = _list_project_folders(root)
    ran = 0
    log.info("Scan start: root=%s folders=%d", root, len(folders))
    for folder in folders:
        try:
            if _process_onedrive_folder(folder):
                ran += 1
        except Exception:
            log.exception("Folder scan failed: %s", folder)
            continue
    log.info("Scan complete: root=%s evaluated=%d", root, ran)
    return ran

def _onedrive_scheduler_loop():
    log = _get_onedrive_logger()
    while True:
        try:
            times = _get_onedrive_scan_times()
            sleep_s = _seconds_until_next_run(times)
            log.info("Scheduler sleeping for %.1f seconds", sleep_s)
            time.sleep(sleep_s)
            if _get_onedrive_scan_enabled():
                log.info("Scheduled scan triggered")
                scan_onedrive_root()
        except Exception:
            # Avoid tight crash loops
            log.exception("Scheduler loop error")
            time.sleep(60)

def start_onedrive_scheduler():
    global _ONEDRIVE_SCHEDULER_STARTED
    if _ONEDRIVE_SCHEDULER_STARTED:
        return
    with _ONEDRIVE_SCHEDULER_LOCK:
        if _ONEDRIVE_SCHEDULER_STARTED:
            return
        t = threading.Thread(target=_onedrive_scheduler_loop, daemon=True)
        t.start()
        _ONEDRIVE_SCHEDULER_STARTED = True

@st.cache_resource
def _start_onedrive_scheduler_once():
    start_onedrive_scheduler()
    return True

# ----------------- UPLOADED FILES MANAGEMENT -----------------
def save_uploaded_file(file_bytes: bytes, filename: str, skip_if_exists: bool = True) -> tuple[str | None, bool]:
    """
    Save uploaded file to uploads folder and database. 
    Returns (saved_file_path or None, is_duplicate_flag).
    If skip_if_exists is True, won't create duplicate DB entries for same file_hash.
    """
    try:
        import hashlib
        import time
        # Generate unique filename: hash + timestamp + original name
        file_hash_full = hashlib.sha256(file_bytes).hexdigest()
        file_hash_short = file_hash_full[:16]
        timestamp = int(time.time())
        safe_name = re.sub(r'[^\w\-_\.]', '_', filename)
        saved_name = f"{file_hash_short}_{timestamp}_{safe_name}"
        saved_path = os.path.join(UPLOADS_DIR, saved_name)
        
        # Save to filesystem (always save physical file for user convenience)
        with open(saved_path, "wb") as f:
            f.write(file_bytes)
        
        # Save to database - check for duplicates first and PREVENT duplicate entries
        is_duplicate = False
        ok, conn, _, _ = _open_mysql_or_create()
        if ok:
            try:
                _ensure_eval_tables(conn)
                file_ext = os.path.splitext(filename)[1].lower().lstrip('.')
                with conn.cursor() as c:
                    # ALWAYS check if file_hash already exists - CRITICAL to prevent duplicates
                    c.execute(
                        "SELECT id, file_path FROM uploaded_rfp_files WHERE file_hash = %s LIMIT 1",
                        (file_hash_full,)
                    )
                    existing = c.fetchone()
                    
                    if existing:
                        # DUPLICATE FOUND - DO NOT INSERT a new row.
                        # If the previously-saved file path is missing, update it to the newly-saved path.
                        is_duplicate = True
                        existing_path = existing.get("file_path")
                        if existing_path and os.path.exists(existing_path):
                            saved_path = existing_path
                        else:
                            # Old path is broken -> update DB to the new saved_path so downstream extraction won't fail.
                            try:
                                c.execute(
                                    "UPDATE uploaded_rfp_files SET file_path=%s, saved_filename=%s, file_size=%s WHERE id=%s",
                                    (saved_path, saved_name, len(file_bytes), existing.get("id")),
                                )
                            except Exception:
                                pass
                        # DO NOT INSERT - prevent duplicate entry
                    elif skip_if_exists:
                        # No duplicate found, insert new record
                        c.execute(
                            """
                            INSERT INTO uploaded_rfp_files (original_filename, saved_filename, file_path, file_size, file_hash, file_type)
                            VALUES (%s, %s, %s, %s, %s, %s)
                            """,
                            (filename, saved_name, saved_path, len(file_bytes), file_hash_full, file_ext)
                        )
                    else:
                        # Legacy behavior - always insert (not recommended)
                        c.execute(
                            """
                            INSERT INTO uploaded_rfp_files (original_filename, saved_filename, file_path, file_size, file_hash, file_type)
                            VALUES (%s, %s, %s, %s, %s, %s)
                            """,
                            (filename, saved_name, saved_path, len(file_bytes), file_hash_full, file_ext)
                        )
                conn.commit()
            except Exception:
                pass
            finally:
                try:
                    conn.close()
                except Exception:
                    pass
        
        return saved_path, is_duplicate
    except Exception as e:
        st.error(f"Failed to save file: {e}")
        return None, False

def list_uploaded_files() -> list[dict]:
    """List all uploaded files from database with metadata."""
    files = []
    ok, conn, _, _ = _open_mysql_or_create()
    if not ok:
        return files
    try:
        _ensure_eval_tables(conn)
        with conn.cursor() as c:
            c.execute("""
                SELECT id, original_filename, saved_filename, file_path, file_size, file_hash, file_type, uploaded_at
                FROM uploaded_rfp_files
                ORDER BY uploaded_at DESC
            """)
            for row in c.fetchall():
                fpath = row.get("file_path")
                # Verify file still exists on disk
                if fpath and os.path.exists(fpath):
                    uploaded_at = row.get("uploaded_at")
                    # Convert datetime to timestamp
                    if uploaded_at:
                        try:
                            if hasattr(uploaded_at, 'timestamp'):
                                mod_ts = uploaded_at.timestamp()
                            else:
                                from datetime import datetime
                                if isinstance(uploaded_at, str):
                                    mod_ts = datetime.fromisoformat(uploaded_at.replace('Z', '+00:00')).timestamp()
                                else:
                                    mod_ts = 0
                        except Exception:
                            mod_ts = 0
                    else:
                        mod_ts = 0
                    files.append({
                        "id": row.get("id"),
                        "filename": row.get("original_filename"),
                        "saved_name": row.get("saved_filename"),
                        "path": fpath,
                        "size": row.get("file_size") or 0,
                        "file_hash": row.get("file_hash"),
                        "file_type": row.get("file_type"),
                        "modified": mod_ts,
                        "uploaded_at": uploaded_at,
                    })
    except Exception:
        pass
    finally:
        try:
            conn.close()
        except Exception:
            pass
    return files

# ----------------- STREAMLIT UI (simplified; no uploaded files listing) -----------------
st.title("📄 AI-Driven RFP Evaluation System")

# Upload RFP - always visible
st.subheader("Upload RFP Document")
uploaded_files = st.file_uploader(
    "Upload RFP Document(s) (PDF/TXT/Image/DOCX/Excel)",
    type=["pdf","txt","docx","jpg","jpeg","png","tiff","tif","xlsx","xls"],
    accept_multiple_files=True,
    key="rfp_uploader_top",
)

def _combine_uploaded_files(files: list, max_chars: int = 250000,
                           prefer_fast_pdf: bool = True,
                           force_pdf_ocr: bool = False,
                           pdf_fast_max_pages: int = 30):
    combined_parts = []
    processed_names = []
    processed_files: list[tuple[str, bytes]] = []
    combined_len = 0
    duplicate_results = []
    duplicates_skipped = 0

    for f in files:
        fname = f.name
        fbytes = f.read()
        try:
            f.seek(0)
        except Exception:
            pass


        # ✅ file-level duplicate detection
        fh = get_file_hash(fbytes)
        is_analyzed, eval_info = check_file_analyzed(fh)
        if is_analyzed and eval_info and eval_info.get("result_text"):
            duplicates_skipped += 1
            duplicate_results.append({"file": fname, "result": eval_info.get("result_text")})
            continue

        # ✅ save file (safe; skip duplicate DB entries)
        save_uploaded_file(fbytes, fname, skip_if_exists=True)

        # ✅ extract
        t = extract_text_any(
            fbytes,
            fname,
            prefer_fast_pdf=bool(prefer_fast_pdf),
            force_pdf_ocr=bool(force_pdf_ocr),
            pdf_fast_max_pages=int(pdf_fast_max_pages),
        ).strip()

        if not t:
            continue

        block = f"\n\n===== FILE: {fname} =====\n{t}"
        if combined_len + len(block) > int(max_chars):
            break
        combined_parts.append(block)
        processed_names.append(fname)
        processed_files.append((fname, fbytes))
        combined_len += len(block)

    combined_text = ("\n".join(combined_parts)).strip()
    return combined_text, processed_names, processed_files, duplicate_results, duplicates_skipped

if uploaded_files and len(uploaded_files) > 0:
    # ------------------------------------------------------------------
    # SINGLE FILE = original behavior (duplicate detection + manual button)
    # MULTI FILE = auto combine + auto evaluate + reuse analyze_with_gpt
    # ------------------------------------------------------------------
    if len(uploaded_files) == 1:
        uploaded_file = uploaded_files[0]
        raw = uploaded_file.read()
        file_hash = get_file_hash(raw)

        # CRITICAL: Check for duplicate FIRST, before any saving or processing
        is_analyzed, eval_info = check_file_analyzed(file_hash)
        force_reanalyze = st.session_state.get("force_reanalyze_top", False)

        if is_analyzed and eval_info and not force_reanalyze:
            created_at = eval_info.get('created_at')
            if created_at:
                if hasattr(created_at, 'strftime'):
                    date_str = created_at.strftime("%Y-%m-%d %H:%M:%S")
                else:
                    date_str = str(created_at)[:19]
            else:
                date_str = "previous analysis"

            st.warning(
                f"⚠️ **DUPLICATE FILE DETECTED** - This file has been analyzed before "
                f"(analyzed on {date_str}). Showing cached results."
            )
            st.success("✅ Previous evaluation results found! Displaying instantly (no re-analysis, no duplicate DB entry).")

            previous_result = eval_info.get("result_text")
            if previous_result:
                st.markdown("---")
                st.markdown("### 📊 Previous Evaluation Results")
                st.markdown(previous_result)
            else:
                st.error("❌ Previous results found but content is empty. Please re-analyze.")

            st.markdown("---")
            col1, col2 = st.columns([1, 4])
            with col1:
                if st.button("🔄 Re-analyze File", key="reanalyze_top", type="primary"):
                    st.session_state["force_reanalyze_top"] = True
                    st.rerun()
            with col2:
                st.caption("Click to run a new analysis (will create new evaluation record)")

        else:
            st.session_state["force_reanalyze_top"] = False

            if not is_analyzed:
                saved_path, is_file_duplicate = save_uploaded_file(raw, uploaded_file.name, skip_if_exists=True)
                if saved_path:
                    if is_file_duplicate:
                        st.warning("⚠️ File hash already exists in database. Skipping duplicate entry.")
                    else:
                        st.info(f"💾 File saved: {os.path.basename(saved_path)}")
            else:
                st.info("📁 File already in database. Proceeding with new analysis...")

            with _st_spinner("🔍 Extracting text via DocTR OCR..."):
                text = extract_text_any(raw, uploaded_file.name)
            if not text.strip():
                _st_safe(st.error, "❌ No text extracted. Try another document.")
                st.stop()
            _st_safe(st.success, "✅ Text extraction complete.")

            if st.button("Run Evaluation", key="run_eval_top"):
                with _st_spinner("🧠 Analyzing RFP using OpenAI (per-question scoring)..."):
                    result, _ = analyze_with_gpt(
                        text,
                        file_name=uploaded_file.name,
                        bypass_decision=bool(st.session_state.get("bypass_decision")),
                        file_hash=file_hash,
                        attached_files=[(uploaded_file.name, raw)],
                    )
                _st_safe(st.markdown, result)

    else:
        # ------------------------------------------------------------------
        # ✅ MULTI-FILE MODE (AUTO): extract + combine + evaluate automatically
        # ------------------------------------------------------------------
        st.info(f"📦 Detected **{len(uploaded_files)} files**. Extracting & combining automatically...")

        max_chars = 250000
        prefer_fast_pdf = True
        force_pdf_ocr = False
        pdf_fast_max_pages = 30

        with _st_spinner("🔍 Extracting + combining all files (duplicates skipped instantly)..."):
            combined_text, processed_names, processed_files, duplicate_results, duplicates_skipped = _combine_uploaded_files(
                uploaded_files,
                max_chars=max_chars,
                prefer_fast_pdf=prefer_fast_pdf,
                force_pdf_ocr=force_pdf_ocr,
                pdf_fast_max_pages=pdf_fast_max_pages,
            )

        # Show cached results instantly for duplicates
        if duplicate_results:
            st.markdown("---")
            st.markdown("## ⚡ Cached Results (Duplicates Skipped)")
            for item in duplicate_results:
                st.markdown(f"### 📌 {item['file']}")
                st.markdown(item["result"])

        # If everything was duplicate
        if duplicates_skipped and not processed_names:
            st.info("✅ All uploaded files were duplicates - no new processing was required.")
            st.stop()

        if not combined_text.strip():
            st.error("❌ Could not extract any new text from the provided files.")
            st.stop()

        st.success(f"✅ Combined text ready from {len(processed_names)} new files.")
        with st.expander("Show processed file list", expanded=False):
            st.write(processed_names)

        # Derive deterministic label from filenames for duplicate detection at batch level
        batch_label = "BATCH_PACKAGE"
        try:
            import hashlib
            names_concat = "||".join(sorted(processed_names))
            h = hashlib.sha256(names_concat.encode("utf-8")).hexdigest()[:16]
            batch_label = f"BATCH_{h}"
        except Exception:
            pass

        # ✅ AUTO RUN evaluation (no extra button)
        with _st_spinner("🧠 Evaluating combined package automatically..."):
            result, _ = analyze_with_gpt(
                combined_text,
                file_name=batch_label,
                bypass_decision=bool(st.session_state.get("bypass_decision")),
                file_hash=get_file_hash(combined_text.encode("utf-8")),
                attached_files=processed_files,
            )

        st.markdown(result)

# ----------------- FOLDER / BATCH EVALUATION -----------------
with st.expander("Folder / Batch Evaluation (multiple documents)", expanded=False):
    st.caption("Upload multiple files OR provide a local folder path. Duplicates are detected instantly and skipped.")

    batch_files = st.file_uploader(
        "Upload multiple documents (PDF/DOCX/TXT/Images/Excel)",
        type=["pdf","txt","docx","jpg","jpeg","png","tiff","tif","xlsx","xls"],
        accept_multiple_files=True,
        key="rfp_uploader_batch",
    )

    folder_path = st.text_input(
        "Or enter folder path (example: C:\\RFPs\\Package_01)",
        value="",
        key="rfp_folder_path",
        help="Reads files from disk on the machine running Streamlit.",
    )

    max_chars = st.number_input(
        "Max combined characters (safety limit)",
        min_value=20000,
        max_value=800000,
        value=250000,
        step=10000,
        key="batch_max_chars",
    )

    prefer_fast_pdf = st.checkbox(
        "Prefer fast PDF text extraction (skip OCR when possible)",
        value=True,
        key="batch_prefer_fast_pdf",
    )
    force_pdf_ocr = st.checkbox(
        "Force OCR for PDFs (slow)",
        value=False,
        key="batch_force_pdf_ocr",
    )
    pdf_fast_max_pages = st.number_input(
        "Fast PDF max pages",
        min_value=1,
        max_value=300,
        value=30,
        step=5,
        key="batch_pdf_fast_max_pages",
    )

    if st.button("Run Folder/Batch Evaluation", key="run_eval_batch"):
        import glob

        # --------------------------------------------------------------
        # Early duplicate detection for batch/folder evaluations. If a
        # folder path is provided, compute a batch label based on the
        # folder name and check the DB for existing evaluations. If
        # found, display the cached result and skip further processing.
        batch_label = None
        if folder_path:
            cleaned_path = folder_path.rstrip("\\/")  # remove trailing slashes/backslashes
            base_name = os.path.basename(cleaned_path)
            batch_label = f"FOLDER_{base_name or 'PACKAGE'}"
            is_done, batch_info = check_batch_analyzed(batch_label)
            if is_done and batch_info and batch_info.get("result_text"):
                st.warning(f"⚠️ **DUPLICATE FOLDER DETECTED** - {base_name or 'PACKAGE'}")
                st.success("✅ Previous evaluation results found! Displaying instantly.")
                st.markdown(batch_info.get("result_text"))
                st.stop()

        files_to_process: list[tuple[str, bytes]] = []

        # Option A: uploaded files
        if batch_files:
            for f in batch_files:
                try:
                    files_to_process.append((f.name, f.read()))
                except Exception:
                    pass

        # Option B: folder path
        if folder_path and os.path.isdir(folder_path):
            exts = ("*.pdf", "*.txt", "*.docx", "*.jpg", "*.jpeg", "*.png", "*.tiff", "*.tif", "*.xlsx", "*.xls")
            disk_files: list[str] = []
            for ex in exts:
                disk_files.extend(glob.glob(os.path.join(folder_path, ex)))
            disk_files = sorted(set(disk_files))
            for fp in disk_files:
                try:
                    with open(fp, "rb") as fh:
                        files_to_process.append((os.path.basename(fp), fh.read()))
                except Exception:
                    pass

        if not files_to_process:
            st.error("❌ No files provided. Upload files or provide a valid folder path.")
            st.stop()

        combined_parts: list[str] = []
        processed_names: list[str] = []
        processed_files: list[tuple[str, bytes]] = []
        combined_len = 0
        n_files = len(files_to_process)

        status_ph = st.empty()
        prog = st.progress(0)

        duplicate_results: list[dict] = []
        duplicates_skipped = 0

        with _st_spinner(f"🔍 Extracting text from {n_files} files (duplicates skipped instantly)..."):
            for i, (fname, fbytes) in enumerate(files_to_process):
                try:
                    prog.progress(min(100, int(100 * ((i + 1) / max(1, n_files)))))
                    status_ph.info(f"Checking {i + 1}/{n_files}: {fname}")

                    if combined_len >= int(max_chars):
                        break

                    # ✅ DUPLICATE DETECTION FIRST (file-level)
                    fh = get_file_hash(fbytes)
                    is_analyzed, eval_info = check_file_analyzed(fh)

                    if is_analyzed and eval_info and eval_info.get("result_text"):
                        duplicates_skipped += 1
                        st.warning(f"⚠️ **DUPLICATE FILE DETECTED** - {fname}")
                        st.success("✅ Previous evaluation results found! Displaying instantly.")
                        duplicate_results.append({"file": fname, "result": eval_info.get("result_text")})
                        continue  # ✅ Skip processing

                    # ✅ Save file (prevents duplicate DB entries)
                    save_uploaded_file(fbytes, fname, skip_if_exists=True)

                    # ✅ Extract text
                    t = extract_text_any(
                        fbytes,
                        fname,
                        prefer_fast_pdf=bool(prefer_fast_pdf),
                        force_pdf_ocr=bool(force_pdf_ocr),
                        pdf_fast_max_pages=int(pdf_fast_max_pages),
                    ).strip()

                    if not t:
                        continue

                    block = f"\n\n===== FILE: {fname} =====\n{t}"
                    combined_parts.append(block)
                    processed_names.append(fname)
                    processed_files.append((fname, fbytes))
                    combined_len += len(block)

                except Exception:
                    continue

        prog.progress(100)
        status_ph.empty()

        # ✅ Show cached results instantly for file-level duplicates
        if duplicate_results:
            st.markdown("---")
            st.markdown("## ⚡ Cached Results (Duplicates Skipped)")
            for item in duplicate_results:
                st.markdown(f"### 📌 {item['file']}")
                st.markdown(item["result"])

        # ✅ If everything was duplicate at file level
        if duplicates_skipped and not processed_names:
            st.info("✅ All uploaded files were duplicates - no new processing was required.")
            st.stop()

        combined_text = ("\n".join(combined_parts)).strip()
        if not combined_text:
            st.error("❌ Could not extract any new text from the provided files.")
            st.stop()

        st.success(f"✅ Combined text ready from {len(processed_names)} new files.")
        with st.expander("Show processed file list", expanded=False):
            st.write(processed_names)

        # ✅ Run evaluation on combined text
        # Reuse batch_label if set earlier (for folder evaluations) to ensure consistent storage in DB.
        # Otherwise derive a deterministic label from uploaded filenames for duplicate detection.
        if not batch_label:
            batch_label = "BATCH_PACKAGE"
            if not folder_path and processed_names:
                try:
                    import hashlib
                    names_concat = "||".join(sorted(processed_names))
                    h = hashlib.sha256(names_concat.encode("utf-8")).hexdigest()[:16]
                    batch_label = f"BATCH_{h}"
                except Exception:
                    pass
        with _st_spinner("🧠 Evaluating combined RFP package..."):
            result, _ = analyze_with_gpt(
                combined_text,
                file_name=batch_label,
                bypass_decision=bool(st.session_state.get("bypass_decision")),
                file_hash=get_file_hash(combined_text.encode("utf-8")),
                attached_files=processed_files,
            )

        st.markdown(result)

# ----------------- OneDrive Auto Sync -----------------
with st.expander("OneDrive Auto Sync", expanded=False):
    st.caption("Uses a locally synced OneDrive/SharePoint folder path. The web link must be synced to this machine.")
    st.text_input(
        "OneDrive root folder (local path)",
        value=_get_onedrive_root(),
        key="onedrive_root_override",
        help="Example: C:\\Users\\<you>\\OneDrive - Ikio\\RFPs",
    )
    st.checkbox(
        "Enable scheduled scan (12:00 AM and 12:00 PM by default)",
        value=_get_onedrive_scan_enabled(),
        key="onedrive_scan_enabled_override",
    )
    st.text_input(
        "Scan times (HH:MM, comma-separated)",
        value=",".join([f"{h:02d}:{m:02d}" for (h, m) in _get_onedrive_scan_times()]),
        key="onedrive_scan_times_override",
    )
    st.number_input(
        "Max combined characters per project",
        min_value=20000,
        max_value=800000,
        value=_get_onedrive_max_chars(),
        step=10000,
        key="onedrive_max_chars_override",
    )
    if st.button("Run OneDrive Scan Now", key="run_onedrive_scan_now"):
        root = _get_onedrive_root()
        if not root or not os.path.isdir(root):
            st.error("❌ OneDrive root folder is not set or not found.")
        else:
            with _st_spinner("Scanning OneDrive folders..."):
                n = scan_onedrive_root()
            st.success(f"✅ OneDrive scan completed. Evaluated {n} folder(s).")

# Start scheduler after UI overrides are applied
if _get_onedrive_scan_enabled():
    _start_onedrive_scheduler_once()

# One-time auto-initialization: ensure website content saved for Ikio Led Lighting LLC, METCO, SUNSPRINT
try:
    if not st.session_state.get("_webctx_init_done"):
        # Skip heavy auto-scrape on first load; user can refresh manually via the UI
        st.session_state["_webctx_init_done"] = True
except Exception:
    pass

col_a, col_b = st.columns([1,3])
with col_a:
    st.selectbox("Model", ["gpt-4o","gpt-4o-mini"], index=0, key="llm_model")
    st.checkbox(
        "Bypass Go/No-Go decision (store full evaluation)",
        value=False,
        key="bypass_decision",
        help="If enabled, the system will not force or apply Go/No-Go decisions. Export JSON will be saved as decision=EVALUATED, comp_name=ALL.",
    )
with col_b:
    with st.expander("Scoring Questions (a..i)", expanded=False):
        st.markdown("\n".join([f"- {ltr}. {txt}" for (ltr, txt) in QUESTIONS_A_I]))

with st.expander("Company Base Locations", expanded=False):
    try:
        rows = []
        db_locs = {}
        ok, conn, _, _ = _open_mysql_or_create()
        if ok:
            try:
                _ensure_company_tables(conn)
                with conn.cursor() as c:
                    c.execute("SELECT company_name, base_location, base_state FROM company_locations")
                    for r in c.fetchall():
                        db_locs[r["company_name"]] = r
            finally:
                try:
                    conn.close()
                except Exception:
                    pass
        
        for comp in COMPANY_LOCATIONS.keys():
            db_row = db_locs.get(comp)
            if db_row:
                loc = db_row.get("base_location") or COMPANY_LOCATIONS.get(comp, "")
                state = db_row.get("base_state") or ""
            else:
                loc = COMPANY_LOCATIONS.get(comp, "")
                # Extract state from default format "City, ST"
                if "," in loc:
                    state = loc.split(",")[-1].strip().upper()
                else:
                    state = ""
            rows.append({"Company": comp, "Base Location": loc or "Not Set", "Base State": state or "Not Set"})
        if rows:
            _st_safe(st.table, rows)
        
        # Edit form
        st.subheader("Edit Company Base Location & State")
        sel_comp = st.selectbox("Select Company", list(COMPANY_LOCATIONS.keys()), key="edit_comp_loc")
        current_row = db_locs.get(sel_comp) if db_locs else None
        col1, col2 = st.columns(2)
        with col1:
            default_loc = (current_row.get("base_location") if current_row and current_row.get("base_location") else COMPANY_LOCATIONS.get(sel_comp, ""))
            new_location = st.text_input("Base Location (e.g., Indianapolis, IN)", 
                                        value=default_loc,
                                        key=f"loc_{sel_comp}")
        with col2:
            # Default state mappings
            default_state_map = {
                "Ikio Led Lighting LLC": "INDIANA",
                "Sunsprint Engineering": "INDIANA",
                "METCO Engineering, Inc.": "TEXAS"
            }
            default_state = (current_row.get("base_state") if current_row and current_row.get("base_state") 
                           else default_state_map.get(sel_comp, ""))
            new_state = st.text_input("Base State (full name, e.g., INDIANA, TEXAS)", 
                                     value=default_state,
                                     key=f"state_{sel_comp}",
                                     help="Enter full state name (e.g., INDIANA, TEXAS, CALIFORNIA)")
        if st.button("Save Location & State", key=f"save_loc_{sel_comp}"):
            ok, conn, _, _ = _open_mysql_or_create()
            if ok:
                try:
                    upsert_company_location(conn, sel_comp, new_location.strip() if new_location else None, new_state.strip().upper() if new_state else None)
                    _st_safe(st.success, f"✅ Saved base location for {sel_comp}")
                except Exception as e:
                    _st_safe(st.error, f"Failed to save: {e}")
                finally:
                    try:
                        conn.close()
                    except Exception:
                        pass
            else:
                _st_safe(st.error, "Database connection failed")
    except Exception as _e:
        _st_safe(st.caption, f"(Could not load DB company locations; showing defaults if available. Error: {_e})")


        

# with st.expander("Company Website Context (Scrape/Refresh)", expanded=False):
#     st.caption("Scrapes Ikio Led Lighting LLC, METCO Engineering, Inc., Sunsprint Engineering websites and stores text in company_web_context.")
#     colr1, colr2, colr3 = st.columns([1,1,2])
#     with colr1:
#         max_pages = st.number_input("Max pages", min_value=1, max_value=200, value=50, step=1)
#     with colr2:
#         max_chars = st.number_input("Max chars", min_value=50000, max_value=1000000, value=300000, step=50000)
#     ignore_ssl = st.checkbox("Ignore SSL errors", value=False)
#     use_sitemap = st.checkbox("Use sitemap boost (if available)", value=True)
#     target_company = st.selectbox("Company (for targeted refresh)", ["All","Ikio Led Lighting LLC","METCO Engineering, Inc.","Sunsprint Engineering"], index=0)
#     if not _SCRAPE_OK:
#         st.error("Website scraping dependencies missing. Install: pip install requests beautifulsoup4")
#     run_refresh = st.button("Refresh & Save Website Context")
#     run_refresh_one = st.button("Refresh Selected Company Only")
#     if run_refresh:
#         with _st_spinner("Scraping websites and saving to DB..."):
#             try:
#                 results = refresh_all_company_web_context(refresh=True, max_pages=int(max_pages), max_chars=int(max_chars), ignore_ssl=bool(ignore_ssl), use_sitemap=bool(use_sitemap))
#                 stats = [{"Company": k, "Chars Saved": len(v or "")} for k, v in results.items()]
#                 _st_safe(st.success, "Saved website context for companies.")
#                 if stats:
#                     _st_safe(st.table, stats)
#                 else:
#                     _st_safe(st.warning, "No website context saved (check connectivity or URLs).")
#             except Exception as e:
#                 _st_safe(st.error, f"Website scraping failed: {e}")
#     if run_refresh_one and target_company != "All":
#         with _st_spinner(f"Refreshing website context for {target_company}..."):
#             try:
#                 res_one = ensure_company_web_context([target_company], refresh=True, max_pages=int(max_pages), max_chars=int(max_chars), ignore_ssl=bool(ignore_ssl), use_sitemap=bool(use_sitemap))
#                 ch = len((res_one.get(target_company) or ""))
#                 if ch > 0:
#                     _st_safe(st.success, f"Saved website context for {target_company} (chars: {ch}).")
#                 else:
#                     _st_safe(st.warning, f"No text extracted for {target_company}. Try increasing pages/chars or enabling SSL ignore.")
#             except Exception as e:
#                 _st_safe(st.error, f"Website scraping failed for {target_company}: {e}")

# with st.expander("Admin: Database Reset (DROP + CREATE)", expanded=False):
#     st.warning("DANGER: Drops and recreates all app tables in the current database.")
#     st.caption(f"Active DB: {DB_CFG.get('database')}")
#     colx1, colx2 = st.columns([1,1])
#     with colx1:
#         confirm = st.text_input("Type RESET to confirm", value="")
#     run_reset = st.button("Drop & Recreate All Tables")
#     if run_reset:
#         if confirm.strip().upper() == "RESET":
#             ok, msg = reset_database_schema(drop_first=True)
#             if ok:
#                 _st_safe(st.success, "All tables dropped and recreated successfully.")
#             else:
#                 _st_safe(st.error, f"Reset failed: {msg}")
#         else:
#             _st_safe(st.error, "Confirmation text mismatch. Type RESET to proceed.")

# ----------------- UPLOADED FILES MANAGEMENT -----------------


def list_uploaded_files() -> list[dict]:
    """List all uploaded files from database with metadata."""
    files = []
    ok, conn, _, _ = _open_mysql_or_create()
    if not ok:
        return files
    try:
        _ensure_eval_tables(conn)
        with conn.cursor() as c:
            c.execute("""
                SELECT id, original_filename, saved_filename, file_path, file_size, file_hash, file_type, uploaded_at
                FROM uploaded_rfp_files
                ORDER BY uploaded_at DESC
            """)
            for row in c.fetchall():
                fpath = row.get("file_path")
                # Verify file still exists on disk
                if fpath and os.path.exists(fpath):
                    uploaded_at = row.get("uploaded_at")
                    # Convert datetime to timestamp
                    if uploaded_at:
                        try:
                            if hasattr(uploaded_at, 'timestamp'):
                                mod_ts = uploaded_at.timestamp()
                            else:
                                from datetime import datetime
                                if isinstance(uploaded_at, str):
                                    mod_ts = datetime.fromisoformat(uploaded_at.replace('Z', '+00:00')).timestamp()
                                else:
                                    mod_ts = 0
                        except Exception:
                            mod_ts = 0
                    else:
                        mod_ts = 0
                    files.append({
                        "id": row.get("id"),
                        "filename": row.get("original_filename"),
                        "saved_name": row.get("saved_filename"),
                        "path": fpath,
                        "size": row.get("file_size") or 0,
                        "file_hash": row.get("file_hash"),
                        "file_type": row.get("file_type"),
                        "modified": mod_ts,
                        "uploaded_at": uploaded_at,
                    })
    except Exception:
        pass
    finally:
        try:
            conn.close()
        except Exception:
            pass
    return files
